import json
import re
from pathlib import Path
from docx import Document


def limpiar_texto(t):
        """Normaliza texto: quita espacios extra y saltos de línea."""
        if not t:
            return ""
        return re.sub(r"\s+", " ", t).strip()

def texto_de_tabla(tabla):
        """Concatena el texto de las primeras filas de una tabla, normalizado."""
        textos = []
        for fila in tabla.rows[:2]:  # solo header y 1ra fila de datos
            for celda in fila.cells:
                textos.append(limpiar_texto(celda.text).lower())
        return " | ".join(textos)


def buscar_tabla(doc, palabras_clave, descripcion=""):
        """
        Busca la primera tabla cuyo contenido inicial incluya TODAS las palabras clave.
        Las palabras se buscan en minúsculas dentro del texto concatenado.
        """
        clave_lower = [p.lower() for p in palabras_clave]
        for i, tabla in enumerate(doc.tables):
            texto = texto_de_tabla(tabla)
            if all(p in texto for p in clave_lower):
                return tabla, i
        raise ValueError(f"No se encontró tabla con {palabras_clave} ({descripcion})")


def buscar_tabla_ra_unidades(doc):
        """
        Busca la tabla de RAs y unidades. Es más flexible que buscar_tabla
        porque acepta variantes del encabezado:
        - "Resultado de aprendizaje y/o desempeños" + "Unidades"
        - "RA" (abreviado) + "Unidades de aprendizaje y contenidos"
        - Variantes con/sin "Indicador de logro" en columna 3
        """
        for i, tabla in enumerate(doc.tables):
            if len(tabla.rows) < 2:
                continue

            # Concatenar texto del header (primera fila)
            celdas_header = [limpiar_texto(c.text).lower() for c in tabla.rows[0].cells]
            texto_header = " | ".join(celdas_header)

            # Debe contener una palabra clave de RA Y una de unidad
            tiene_ra = any(
                "resultado de aprendizaje" in c
                or "resultados de aprendizaje" in c
                or c.strip() == "ra"
                or c.startswith("ra ")
                for c in celdas_header
            )
            tiene_unidad = any("unidad" in c for c in celdas_header)

            if tiene_ra and tiene_unidad:
                return tabla, i

        raise ValueError("No se encontró tabla de RAs y unidades en este documento")

def extraer_identificacion(doc):
        """Extrae los datos básicos de la tabla 0, tolerando variantes de etiqueta."""
        tabla = doc.tables[0]

        # Mapeo de etiquetas variantes a un nombre canónico
        ALIAS = {
            "facultad": "facultad",
            "carrera": "carrera",
            "carreras": "carrera",
            "nombre": "nombre",
            "codigo": "codigo",
            "código": "codigo",
            "codigos": "codigo",
            "códigos": "codigo",
            "nivel": "nivel",
            "duracion": "duracion",
            "duración": "duracion",
            "requisito(s)": "requisitos",
            "requisitos": "requisitos",
        }

        def normalizar_clave(texto):
            """Quita los dos puntos finales y baja a minúsculas para comparar."""
            return texto.rstrip(":").strip().lower()

        campos = {}
        for fila in tabla.rows:
            celdas = [limpiar_texto(c.text) for c in fila.cells]
            # Recorrer todas las celdas buscando pares "etiqueta : valor"
            i = 0
            while i < len(celdas) - 1:
                clave_raw = celdas[i]
                valor = celdas[i + 1]
                if clave_raw and valor and clave_raw.endswith(":"):
                    clave_norm = normalizar_clave(clave_raw)
                    if clave_norm in ALIAS:
                        campos[ALIAS[clave_norm]] = valor
                        i += 2
                        continue
                i += 1
        return campos

def extraer_responsables(doc):
        """Busca la tabla de responsables/docente/versión por su contenido."""
        try:
            tabla, _ = buscar_tabla(
                doc,
                palabras_clave=["responsable", "programa"],
                descripcion="tabla de responsables"
            )
        except ValueError:
            return {}

        ALIAS = {
            "responsable(s) del programa": "responsable",
            "responsables del programa": "responsable",
            "docente(s) a cargo": "docente_a_cargo",
            "docentes a cargo": "docente_a_cargo",
            "versión / fecha de actualización": "version",
            "version / fecha de actualizacion": "version",
            "versión": "version",
            "version": "version",
        }

        datos = {}
        for fila in tabla.rows:
            celdas = [limpiar_texto(c.text) for c in fila.cells]
            if len(celdas) >= 2 and celdas[0]:
                clave = celdas[0].rstrip(":").strip().lower()
                if clave in ALIAS:
                    datos[ALIAS[clave]] = celdas[1]
        return datos


def extraer_metodologias(doc):
        """
        Localiza la tabla de metodologías entre la de RAs y la de evaluaciones.
        Si existe, separa por puntos/punto-y-coma.
        """
        try:
            _, idx_unidades = buscar_tabla_ra_unidades(doc)
        except ValueError:
            return []

        # Buscamos la siguiente tabla que contenga palabras de metodología/enseñanza
        palabras_metod = [
            "clases", "taller", "metodolog", "enseñ", "aprendizaje",
            "expositivas", "resolución de problemas", "estudio"
        ]

        for i in range(idx_unidades + 1, min(idx_unidades + 4, len(doc.tables))):
            tabla = doc.tables[i]
            texto = " ".join(
                limpiar_texto(c.text)
                for fila in tabla.rows
                for c in fila.cells
            )
            texto_lower = texto.lower()

            # Saltar si parece bibliografía o evaluación
            if "bibliografía" in texto_lower or "porcentaje" in texto_lower:
                continue

            # Aceptar si menciona palabras de metodología
            if any(p in texto_lower for p in palabras_metod):
                # Limpiar prefijos
                if ":" in texto[:50]:
                    texto = texto.split(":", 1)[1].strip()
                items = [m.strip() for m in re.split(r"[.;]", texto) if m.strip() and len(m.strip()) > 10]
                return items

        return []


def extraer_evaluaciones(doc):
        """
        Busca la tabla de evaluaciones.
        Formato 1: tabla con header 'Tipo de evaluación' + 'Porcentaje' → estructurada
        Formato 2: tabla narrativa con texto descriptivo → guardamos texto completo
        """
        # Intento 1: formato estructurado
        try:
            tabla, _ = buscar_tabla(
                doc,
                palabras_clave=["tipo de evaluación", "porcentaje"],
                descripcion="tabla de evaluaciones estructurada"
            )
            evaluaciones = []
            for fila in tabla.rows[1:]:
                celdas = [limpiar_texto(c.text) for c in fila.cells]
                if len(celdas) < 2 or not celdas[0]:
                    continue
                porcentaje = celdas[1]
                if len(porcentaje) > 100:
                    continue
                if "%" not in porcentaje:
                    continue
                evaluaciones.append({
                    "tipo": celdas[0],
                    "porcentaje": porcentaje
                })
            if evaluaciones:
                return evaluaciones
        except ValueError:
            pass

        # Intento 2: tabla narrativa con palabras clave de evaluación
        palabras_eval = ["evaluación", "evaluacion", "evaluaciones", "prueba", "examen"]
        for tabla in doc.tables:
            texto = " ".join(
                limpiar_texto(c.text).lower()
                for fila in tabla.rows
                for c in fila.cells
            )
            # Buscamos texto donde se mencione evaluación pero no sea bibliografía
            if any(p in texto for p in palabras_eval) and "bibliografía" not in texto:
                # Tomamos solo si el texto es razonablemente corto (no es el programa entero)
                if 100 < len(texto) < 3000:
                    return [{"tipo": "descripcion_narrativa", "porcentaje": "", "texto": texto}]

        return []


        """
        Busca la tabla que contiene RAs y unidades. Tolera variantes:
        - 2 columnas (RAs+Unidad, Indicador opcional) o 3 columnas
        - RAs con o sin nivel de dominio
        """
        tabla, idx = buscar_tabla(
            doc,
            palabras_clave=["resultado de aprendizaje", "unidades"],
            descripcion="tabla de RAs y unidades"
        )

        n_cols = len(tabla.rows[0].cells)
        filas = []

        for fila in tabla.rows[1:]:  # saltamos header
            celdas = [limpiar_texto(c.text) for c in fila.cells]
            if not celdas or not celdas[0]:
                continue


def extraer_ra_unidades(doc):
        """
        Busca la tabla que contiene RAs y unidades. Tolera variantes:
        - 2 columnas (RAs+Unidad, Indicador opcional) o 3 columnas
        - RAs con o sin nivel de dominio
        - Encabezados que digan 'RA' o 'Resultado de aprendizaje'
        """
        try:
            tabla, idx = buscar_tabla_ra_unidades(doc)
        except ValueError:
            return []  # programa sin RAs declarados (electivos, idiomas, TIPE...)

        n_cols = len(tabla.rows[0].cells)
        filas = []

        for fila in tabla.rows[1:]:  # saltamos header
            celdas = [limpiar_texto(c.text) for c in fila.cells]
            if not celdas or not celdas[0]:
                continue

            ras_texto = celdas[0]
            unidad_texto = celdas[1] if len(celdas) > 1 else ""
            logro_texto = celdas[2] if n_cols >= 3 and len(celdas) > 2 else ""

            # Parsear códigos de RA tolerando varios formatos:
            # "CL2, N2, RA1." o "CL1, RA.1:" o "CL1,N2,RA1;"
            patron = r"\b([A-Z]{1,3}\d+(?:,\s*N\d+)?,\s*RA\.?\d+)"
            codigos_ra = re.findall(patron, ras_texto)
            codigos_ra = [
                re.sub(r"RA\.(\d+)", r"RA\1", c).strip()
                for c in codigos_ra
            ]

            filas.append({
                "ras": codigos_ra,
                "unidad_y_contenidos": unidad_texto,
                "indicador_logro": logro_texto
            })

        return filas

def extraer_programa(ruta):
        doc = Document(ruta)
        return {
            "archivo": Path(ruta).name,
            "identificacion": extraer_identificacion(doc),
            "responsables": extraer_responsables(doc),
            "metodologias": extraer_metodologias(doc),
            "evaluaciones": extraer_evaluaciones(doc),
            "unidades": extraer_ra_unidades(doc)
        }

if __name__ == "__main__":
        ruta = "data/programas/Semestre 5/IMAT321 Análisis.docx"
        programa = extraer_programa(ruta)

        # Resumen en consola
        ident = programa["identificacion"]
        print(f"Asignatura extraída: {ident.get('codigo', '?')} - {ident.get('nombre', '?')}")
        print(f"  Facultad: {ident.get('facultad', '?')}")
        print(f"  Carrera: {ident.get('carrera', '?')}")
        print(f"  Nivel: {ident.get('nivel', '?')}")
        print(f"  Docente a cargo: {programa['responsables'].get('docente_a_cargo', '?')}")
        print(f"  Responsable: {programa['responsables'].get('responsable', '?')}")
        print(f"  Versión: {programa['responsables'].get('version', '?')}")
        print(f"  Metodologías: {len(programa['metodologias'])}")
        for m in programa['metodologias']:
            print(f"    · {m[:70]}")
        print(f"  Evaluaciones: {len(programa['evaluaciones'])}")
        for e in programa['evaluaciones']:
            print(f"    · {e['tipo'][:50]} → {e['porcentaje']}")
        print(f"  Unidades detectadas: {len(programa['unidades'])}")
        for i, u in enumerate(programa['unidades'], start=1):
            print(f"    Unidad {i}: {len(u['ras'])} RAs · '{u['unidad_y_contenidos'][:50]}...'")

        # Guardar como JSON
        salida = Path("data/programa_ejemplo.json")
        salida.write_text(
            json.dumps(programa, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        print(f"\nGuardado en: {salida}")