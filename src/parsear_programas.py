"""
parsear_programas.py — Parser de programas .docx para el sistema de gestión curricular ICM.
Extrae toda la información relevante de un documento Word de programa de asignatura.
"""
import re
from pathlib import Path
from docx import Document

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _cell_text(cell):
    """Obtiene el texto completo de una celda."""
    return cell.text.strip()


def _tc_text(tc_elem):
    """Obtiene texto de un elemento tc XML."""
    return "".join(t.text for t in tc_elem.iter(f"{{{WNS}}}t") if t.text).strip()


def _tabla_texto(tabla):
    """Devuelve el texto completo de una tabla."""
    return "\n".join(
        " | ".join(c.text.strip() for c in row.cells)
        for row in tabla.rows
    )


def _inferir_semestre_desde_ruta(ruta):
    """Infiere el semestre desde el path del archivo (ej: 'Semestre 3/')."""
    ruta_str = str(ruta)
    m = re.search(r'Semestre\s+(\d+)', ruta_str, re.IGNORECASE)
    if m:
        return int(m.group(1))
    return None


def _inferir_semestre_desde_nivel(nivel_str):
    """
    Infiere el semestre desde el texto de nivel.
    Ej: '5to Semestre', 'I Semestre del 1° Año' = 1
    """
    if not nivel_str:
        return None
    # Buscar número arábigo
    m = re.search(r'(\d+)[°º]?\s*[Ss]emestre', nivel_str)
    if m:
        return int(m.group(1))
    # Buscar romano
    romanos = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5,
               "VI": 6, "VII": 7, "VIII": 8, "IX": 9, "X": 10}
    m = re.search(r'\b(I{1,3}V?|VI{0,3}|IX|X)\b\s*[Ss]emestre', nivel_str)
    if m:
        return romanos.get(m.group(1).upper())
    return None


def _parsear_tabla_identificacion(tabla):
    """
    Parsea la tabla de identificación (Table 0, ~8 filas x 6 cols).
    Retorna dict con identificacion + horas.
    """
    result = {
        "codigo": "", "nombre": "", "nivel": "", "duracion": "",
        "facultad": "", "carrera": "", "requisitos": "",
        "horas_directa": None, "horas_autonoma": None,
        "semanas": None, "creditos": None
    }

    rows = tabla.rows
    for i, row in enumerate(rows):
        cells = [_cell_text(c) for c in row.cells]
        # Deduplicar celdas combinadas (python-docx repite el texto en celdas fusionadas)
        cells_unique = []
        seen = set()
        for c in cells:
            if c not in seen:
                cells_unique.append(c)
                seen.add(c)
            else:
                cells_unique.append("")

        cells_str = " ".join(cells_unique).lower()

        # Fila con Facultad / Carreras
        if "facultad" in cells_str or "facultades" in cells_str:
            # Buscar valor en celdas posteriores
            for j, c in enumerate(cells_unique):
                if "facultad" in c.lower() and j + 1 < len(cells_unique):
                    result["facultad"] = cells_unique[j + 1].strip()
                if "carrera" in c.lower() and j + 1 < len(cells_unique):
                    result["carrera"] = cells_unique[j + 1].strip()

        # Fila con Nombre / Código
        elif "nombre" in cells_str and ("código" in cells_str or "codigo" in cells_str):
            for j, c in enumerate(cells_unique):
                if "nombre" in c.lower() and j + 1 < len(cells_unique):
                    result["nombre"] = cells_unique[j + 1].strip()
                if ("código" in c.lower() or "codigo" in c.lower()) and j + 1 < len(cells_unique):
                    result["codigo"] = cells_unique[j + 1].strip()

        # Fila con Nivel / Duración
        elif "nivel" in cells_str and ("duración" in cells_str or "duracion" in cells_str):
            for j, c in enumerate(cells_unique):
                if "nivel" in c.lower() and j + 1 < len(cells_unique):
                    result["nivel"] = cells_unique[j + 1].strip()
                if ("duración" in c.lower() or "duracion" in c.lower()) and j + 1 < len(cells_unique):
                    result["duracion"] = cells_unique[j + 1].strip()

        # Fila con Requisito
        elif "requisito" in cells_str:
            vals = [c for c in cells_unique if c and "requisito" not in c.lower()]
            result["requisitos"] = ", ".join(vals) if vals else ""

        # Fila de valores de horas (contiene números, la última fila de datos)
        # La fila (A)(B)(C)(D)(E)(F) es la de etiquetas, la siguiente son los valores
        elif re.search(r'\(A\)', cells_str) or re.search(r'\(B\)', cells_str):
            # Esta es la fila de etiquetas, la siguiente tiene los valores
            if i + 1 < len(rows):
                val_cells = [_cell_text(c) for c in rows[i + 1].cells]
                # Deduplicar
                val_unique = []
                seen2 = set()
                for c in val_cells:
                    if c not in seen2:
                        val_unique.append(c)
                        seen2.add(c)
                    else:
                        val_unique.append("")
                vals = [v for v in val_unique if v.strip()]

                def to_float(s):
                    try:
                        return float(s.replace(",", "."))
                    except:
                        return None

                def to_int(s):
                    try:
                        return int(float(s.replace(",", ".")))
                    except:
                        return None

                if len(vals) >= 1:
                    result["horas_directa"] = to_float(vals[0])
                if len(vals) >= 2:
                    result["horas_autonoma"] = to_float(vals[1])
                if len(vals) >= 4:
                    result["semanas"] = to_int(vals[3])
                if len(vals) >= 6:
                    result["creditos"] = to_int(vals[5])

    return result


def _es_tabla_unidades(tabla):
    """Detecta si una tabla es la tabla de Unidades (3 cols con cabecera de RA)."""
    if len(tabla.columns) < 2:
        return False
    if not tabla.rows:
        return False
    first_row_text = " ".join(_cell_text(c).lower() for c in tabla.rows[0].cells)
    return ("resultado de aprendizaje" in first_row_text or
            "unidades de aprendizaje" in first_row_text)


def _parsear_unidades(tabla):
    """Parsea la tabla de unidades. Retorna lista de dicts."""
    unidades = []
    rows = tabla.rows
    if len(rows) < 2:
        return unidades

    n_cols = len(tabla.columns)

    for row in rows[1:]:
        cells = [_cell_text(c) for c in row.cells]
        if not any(cells):
            continue

        # Deduplicar celdas combinadas
        cells_unique = []
        seen = set()
        for c in cells:
            if c not in seen or not c:
                cells_unique.append(c)
                seen.add(c)
            else:
                cells_unique.append("")

        ra_col = cells_unique[0] if len(cells_unique) > 0 else ""
        contenido_col = cells_unique[1] if len(cells_unique) > 1 else ""
        indicador_col = cells_unique[2] if len(cells_unique) > 2 else ""

        # Extraer códigos de RA (separados por ; o salto de línea)
        ra_codes = []
        if ra_col:
            for part in re.split(r'[;\n]', ra_col):
                part = part.strip()
                if part and re.match(r'[A-Z]{2}\d', part):
                    ra_codes.append(part)

        if contenido_col or ra_codes:
            unidades.append({
                "ra_codes": ra_codes,
                "titulo": contenido_col[:100] if contenido_col else "",
                "contenidos": contenido_col,
                "indicador_logro": indicador_col
            })

    return unidades


def _parsear_ra_generales(doc, idx_tabla_unidades):
    """
    Busca los RA de la sección 'RESULTADOS DE APRENDIZAJE Y DESEMPEÑOS'
    en párrafos o tabla antes de las unidades.
    """
    ra_codes = []
    body = doc.element.body
    elems = list(body)

    # Buscar el marcador de sección RA
    for i, elem in enumerate(elems):
        tag = elem.tag.split("}")[-1]
        if tag == "p":
            text = "".join(t.text for t in elem.iter(f"{{{WNS}}}t") if t.text).strip()
            if "RESULTADO" in text.upper() and "APRENDIZAJE" in text.upper():
                # Buscar párrafos con códigos RA entre aquí y UNIDADES
                for j in range(i+1, min(i+50, len(elems))):
                    e = elems[j]
                    etag = e.tag.split("}")[-1]
                    if etag == "p":
                        txt = "".join(t.text for t in e.iter(f"{{{WNS}}}t") if t.text).strip()
                        if "UNIDAD" in txt.upper() and "APRENDIZAJE" in txt.upper():
                            break
                        # RA codes tienen formato CL1, N1, RA1 o similar
                        m = re.match(r'([A-Z]{2}\d[\w,\s\.]*RA[\.\d]+)', txt)
                        if m:
                            ra_codes.append(txt.split(":")[0].strip())
                break

    return ra_codes


def _es_tabla_metodologia_checkbox(tabla):
    """Detecta si la tabla tiene checkboxes de metodología."""
    full_text = _tabla_texto(tabla).lower()
    return any(m in full_text for m in ["clase expositiva", "aprendizaje basado", "seminario", "taller", "laboratorio"])


def _parsear_metodologias_checkbox(tabla):
    """Extrae metodologías marcadas con X en una tabla de checkboxes."""
    metodologias = []
    for row in tabla.rows:
        cells = [_cell_text(c) for c in row.cells]
        for i in range(len(cells)):
            if cells[i].strip().upper() in ("X", "✓", "✗", "☑") and i + 1 < len(cells):
                metodologias.append(cells[i + 1].strip())
            elif i > 0 and cells[i-1].strip().upper() in ("X", "✓") and cells[i]:
                pass  # ya capturado
    return metodologias


def _es_tabla_evaluacion(tabla):
    """Detecta si la tabla es de evaluaciones."""
    if not tabla.rows:
        return False
    first_row = " ".join(_cell_text(c).lower() for c in tabla.rows[0].cells)
    return ("evaluaci" in first_row or "ponderaci" in first_row or "porcentaje" in first_row) and \
           len(tabla.columns) >= 2


def _parsear_evaluaciones(tabla):
    """Parsea la tabla de evaluaciones."""
    evaluaciones = []
    for row in tabla.rows[1:]:
        cells = [_cell_text(c) for c in row.cells]
        tipo = cells[0].strip() if cells else ""
        porc = cells[1].strip() if len(cells) > 1 else ""
        if tipo and "tipo" not in tipo.lower():
            evaluaciones.append({"tipo": tipo, "porcentaje": porc})
    return evaluaciones


def _es_tabla_bibliografia(tabla):
    """Detecta si la tabla es de bibliografía."""
    if not tabla.rows:
        return False
    full_text = " ".join(_cell_text(c).lower() for c in tabla.rows[0].cells)
    return "autor" in full_text or "título" in full_text or "titulo" in full_text or \
           "editorial" in full_text or "isbn" in full_text


def _parsear_bibliografia(tabla, tipo="basica"):
    """Parsea una tabla de bibliografía. Retorna lista de dicts."""
    entradas = []
    rows = tabla.rows
    if len(rows) < 2:
        return entradas

    for row in rows[1:]:
        cells = [_cell_text(c) for c in row.cells]
        if len(cells) >= 3 and any(cells):
            # N° | Autor | Título | Editorial | Año | ISBN | Ejemplares
            entradas.append({
                "numero": cells[0] if len(cells) > 0 else "",
                "autor": cells[1] if len(cells) > 1 else "",
                "titulo": cells[2] if len(cells) > 2 else "",
                "editorial": cells[3] if len(cells) > 3 else "",
                "anio": cells[4] if len(cells) > 4 else "",
                "isbn": cells[5] if len(cells) > 5 else "",
                "ejemplares": cells[6] if len(cells) > 6 else "",
            })
    return entradas


def _parsear_linkografia(tabla):
    """Parsea la tabla de linkografía."""
    entradas = []
    rows = tabla.rows
    if len(rows) < 2:
        return entradas

    for row in rows[1:]:
        cells = [_cell_text(c) for c in row.cells]
        if len(cells) >= 2 and any(cells):
            entradas.append({
                "tipo_doc": cells[0] if len(cells) > 0 else "",
                "autor": cells[1] if len(cells) > 1 else "",
                "titulo_articulo": cells[2] if len(cells) > 2 else "",
                "anio": cells[3] if len(cells) > 3 else "",
                "titulo_revista": cells[4] if len(cells) > 4 else "",
                "volumen": cells[5] if len(cells) > 5 else "",
                "url": cells[6] if len(cells) > 6 else "",
                "disponible": cells[7] if len(cells) > 7 else "",
            })
    return entradas


def _encontrar_texto_entre_secciones(doc, inicio_marker, fin_marker=None):
    """Extrae el texto entre dos marcadores de sección."""
    parrafos = doc.paragraphs
    textos = []
    capturando = False

    for p in parrafos:
        texto = p.text.strip()
        if inicio_marker.lower() in texto.lower():
            capturando = True
            continue
        if fin_marker and fin_marker.lower() in texto.lower():
            break
        if capturando and texto:
            textos.append(texto)

    return "\n".join(textos)


def _get_body_elements(doc):
    return list(doc.element.body)


def _get_paragraph_text(elem):
    return "".join(t.text for t in elem.iter(f"{{{WNS}}}t") if t.text).strip()


def parsear_docx(ruta_docx: str) -> dict:
    """
    Parsea un archivo .docx de programa de asignatura.
    Retorna un diccionario con toda la información extraída.
    """
    ruta = Path(ruta_docx)
    doc = Document(str(ruta))

    resultado = {
        "archivo": str(ruta),
        "identificacion": {
            "codigo": "", "nombre": "", "nivel": "", "duracion": "",
            "facultad": "", "carrera": "", "requisitos": "",
            "horas_directa": None, "horas_autonoma": None,
            "semanas": None, "creditos": None
        },
        "descripcion": "",
        "aporte_perfil": "",
        "responsables": {"responsable": "", "docente_a_cargo": "", "version": ""},
        "ra_codes": [],
        "unidades": [],
        "metodologias": [],
        "evaluaciones": [],
        "bibliografia_basica": [],
        "bibliografia_complementaria": [],
        "linkografia": [],
        "otros_recursos": "",
        "semestre": None
    }

    tablas = doc.tables

    # ── 1. Tabla de identificación (Table 0) ──────────────────────
    if tablas:
        ident = _parsear_tabla_identificacion(tablas[0])
        resultado["identificacion"].update(ident)

    # ── 2. Inferir semestre ────────────────────────────────────────
    semestre = _inferir_semestre_desde_ruta(ruta)
    if semestre is None:
        semestre = _inferir_semestre_desde_nivel(resultado["identificacion"].get("nivel", ""))
    resultado["semestre"] = semestre

    # ── 3. Leer cuerpo del documento por secciones ─────────────────
    body_elems = _get_body_elements(doc)

    # Mapear índices de párrafos y tablas
    seccion_actual = None
    secciones_idx = {}
    tablas_por_seccion = {}
    tabla_counter = 0
    tabla_idx_map = {}  # tabla_counter -> tabla object

    for elem in body_elems:
        tag = elem.tag.split("}")[-1]
        if tag == "p":
            txt = _get_paragraph_text(elem)
            TU = txt.upper()
            if "DESCRIPCIÓN DE LA ASIGNATURA" in TU:
                seccion_actual = "descripcion"
                secciones_idx[seccion_actual] = []
            elif "APORTE AL PERFIL" in TU:
                seccion_actual = "aporte_perfil"
                secciones_idx[seccion_actual] = []
            elif "RESULTADO" in TU and "APRENDIZAJE" in TU and "DESEMPEÑO" in TU:
                seccion_actual = "ras"
                secciones_idx[seccion_actual] = []
            elif "RESULTADO" in TU and "APRENDIZAJE" in TU and seccion_actual != "ras":
                seccion_actual = "ras"
                secciones_idx.setdefault(seccion_actual, [])
            elif "UNIDADES DE APRENDIZAJE" in TU and "CONTENIDO" in TU:
                seccion_actual = "unidades"
                secciones_idx[seccion_actual] = []
            elif "ESTRATEGIA" in TU and ("ENSEÑANZA" in TU or "METODOLOG" in TU):
                seccion_actual = "metodologia"
                secciones_idx[seccion_actual] = []
            elif "METODOLOG" in TU and "EVALUACI" in TU:
                seccion_actual = "evaluacion"
                secciones_idx[seccion_actual] = []
            elif "BIBLIOGRAFÍA" in TU or "BIBLIOGRAFIA" in TU:
                if "BÁSICA" in TU or "BASICA" in TU or "OBLIGATORIA" in TU:
                    seccion_actual = "biblio_basica"
                elif "COMPLEMENTARIA" in TU:
                    seccion_actual = "biblio_complementaria"
                else:
                    seccion_actual = "biblio_basica"
                secciones_idx.setdefault(seccion_actual, [])
            elif "LINKOGRAFÍA" in TU or "LINKOGRAFIA" in TU:
                seccion_actual = "linkografia"
                secciones_idx[seccion_actual] = []
            elif "OTROS RECURSOS" in TU:
                seccion_actual = "otros_recursos"
                secciones_idx[seccion_actual] = []
            elif "DATOS ACTUALIZACIÓN" in TU or "DATOS ACTUALIZACION" in TU:
                seccion_actual = "datos_actualizacion"
                secciones_idx.setdefault(seccion_actual, [])
            elif seccion_actual and txt:
                if seccion_actual in secciones_idx:
                    secciones_idx[seccion_actual].append(txt)
        elif tag == "tbl":
            tabla_idx_map[tabla_counter] = seccion_actual
            tabla_counter += 1

    # ── 4. Descripción ─────────────────────────────────────────────
    desc_textos = secciones_idx.get("descripcion", [])
    resultado["descripcion"] = "\n".join(desc_textos).strip()

    # Si está vacío, buscar en tabla de celda única
    if not resultado["descripcion"]:
        for t in tablas[1:6]:
            if len(t.rows) == 1 and len(t.columns) == 1:
                txt = _cell_text(t.rows[0].cells[0])
                if len(txt) > 50:
                    resultado["descripcion"] = txt
                    break

    # ── 5. Aporte al perfil ────────────────────────────────────────
    aporte_textos = secciones_idx.get("aporte_perfil", [])
    resultado["aporte_perfil"] = "\n".join(aporte_textos).strip()

    # Si no se encontró, buscar en tablas
    if not resultado["aporte_perfil"]:
        for i, t in enumerate(tablas[1:8], 1):
            if len(t.rows) == 1 and len(t.columns) == 1:
                txt = _cell_text(t.rows[0].cells[0])
                if len(txt) > 30 and not resultado["descripcion"].startswith(txt[:20]):
                    if resultado["descripcion"] and len(resultado["aporte_perfil"]) == 0:
                        resultado["aporte_perfil"] = txt
                        break

    # ── 6. RAs generales ──────────────────────────────────────────
    ra_textos = secciones_idx.get("ras", [])
    for txt in ra_textos:
        m = re.match(r'([A-Z]{2}\d[\w,\s\.]*)', txt.strip())
        if m and re.search(r'RA[\.\d]', txt):
            cod = txt.split(":")[0].strip()
            resultado["ra_codes"].append(cod)

    # ── 7. Unidades ────────────────────────────────────────────────
    for ti, t in enumerate(tablas):
        if _es_tabla_unidades(t):
            resultado["unidades"] = _parsear_unidades(t)
            break

    # ── 8. Metodologías ───────────────────────────────────────────
    metod_section = secciones_idx.get("metodologia", [])

    # Buscar tabla de metodología
    metod_tabla_idx = None
    for ti, seccion in tabla_idx_map.items():
        if seccion == "metodologia" and ti < len(tablas):
            metod_tabla_idx = ti
            break

    if metod_tabla_idx is not None and metod_tabla_idx < len(tablas):
        t = tablas[metod_tabla_idx]
        if _es_tabla_metodologia_checkbox(t):
            resultado["metodologias"] = _parsear_metodologias_checkbox(t)
        else:
            # Texto libre en la tabla
            texto = _tabla_texto(t)
            if texto.strip():
                resultado["metodologias"] = [texto.strip()]

    if not resultado["metodologias"] and metod_section:
        resultado["metodologias"] = metod_section

    # ── 9. Evaluaciones ───────────────────────────────────────────
    for ti, seccion in tabla_idx_map.items():
        if seccion == "evaluacion" and ti < len(tablas):
            t = tablas[ti]
            if _es_tabla_evaluacion(t):
                resultado["evaluaciones"] = _parsear_evaluaciones(t)
                break

    # ── 10. Bibliografía ──────────────────────────────────────────
    for ti, seccion in tabla_idx_map.items():
        if ti >= len(tablas):
            continue
        t = tablas[ti]
        if seccion == "biblio_basica" and _es_tabla_bibliografia(t):
            entradas = _parsear_bibliografia(t, "basica")
            resultado["bibliografia_basica"].extend(entradas)
        elif seccion == "biblio_complementaria" and _es_tabla_bibliografia(t):
            entradas = _parsear_bibliografia(t, "complementaria")
            resultado["bibliografia_complementaria"].extend(entradas)

    # Fallback: buscar tablas de bibliografía por contenido
    if not resultado["bibliografia_basica"] and not resultado["bibliografia_complementaria"]:
        for ti, t in enumerate(tablas):
            if not t.rows:
                continue
            header = " ".join(_cell_text(c).lower() for c in t.rows[0].cells)
            if "autor" in header and ("título" in header or "titulo" in header or "isbn" in header):
                if ti < len(tablas) - 2 and not resultado["bibliografia_basica"]:
                    resultado["bibliografia_basica"] = _parsear_bibliografia(t)
                elif not resultado["bibliografia_complementaria"]:
                    resultado["bibliografia_complementaria"] = _parsear_bibliografia(t)

    # ── 11. Linkografía ───────────────────────────────────────────
    for ti, seccion in tabla_idx_map.items():
        if seccion == "linkografia" and ti < len(tablas):
            t = tablas[ti]
            resultado["linkografia"] = _parsear_linkografia(t)
            break

    # ── 12. Otros recursos ────────────────────────────────────────
    otros_textos = secciones_idx.get("otros_recursos", [])
    resultado["otros_recursos"] = "\n".join(otros_textos).strip()

    # Buscar en tabla también
    if not resultado["otros_recursos"]:
        for ti, seccion in tabla_idx_map.items():
            if seccion == "otros_recursos" and ti < len(tablas):
                t = tablas[ti]
                resultado["otros_recursos"] = _tabla_texto(t).strip()
                break

    # ── 13. Datos de actualización ────────────────────────────────
    for ti, seccion in tabla_idx_map.items():
        if seccion == "datos_actualizacion" and ti < len(tablas):
            t = tablas[ti]
            for row in t.rows:
                cells = [_cell_text(c) for c in row.cells]
                if len(cells) >= 2:
                    label = cells[0].lower()
                    val = cells[1]
                    if "responsable" in label:
                        resultado["responsables"]["responsable"] = val
                    elif "docente" in label:
                        resultado["responsables"]["docente_a_cargo"] = val
                    elif "versión" in label or "version" in label or "fecha" in label:
                        resultado["responsables"]["version"] = val
            break

    return resultado


if __name__ == "__main__":
    import sys, json
    if len(sys.argv) < 2:
        print("Uso: python3 src/parsear_programas.py <ruta.docx>")
        sys.exit(1)
    resultado = parsear_docx(sys.argv[1])
    print(json.dumps(resultado, ensure_ascii=False, indent=2))
