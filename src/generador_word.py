"""
generador_word.py — Genera programa de asignatura fiel al documento original.

`generar_programa_individual(asignatura_id)` construye el documento Word desde
la BD con TODAS las secciones. `generar_mapa_progreso()` genera el mapa de
progreso global.
"""
import sqlite3, re, copy, os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DB_PATH  = Path("data/sistema.db")
OUT_DIR  = Path("data/output")
PROG_DIR = Path("data/programas")

# Alias usados por la nueva implementación
RUTA_DB = DB_PATH
RUTA_OUTPUT = OUT_DIR
LOGO_PATH = Path("data/uv_logo_nuevo.png")

UV_BLUE = RGBColor(0x1F, 0x4E, 0x79)
UV_LIGHT = RGBColor(0xBD, 0xD7, 0xEE)
VERDE = RGBColor(0x37, 0x56, 0x23)

TIPO_COLOR = {
    "licenciatura": UV_BLUE,
    "titulo": VERDE,
    "sello_uv": RGBColor(0x84, 0x3C, 0x0C),
}

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── BD ────────────────────────────────────────────────────────────

def conectar():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn

def get_programa(asig_id):
    conn = conectar()
    asig = dict(conn.execute(
        "SELECT * FROM asignaturas WHERE id=?", (asig_id,)
    ).fetchone())
    unidades = [dict(r) for r in conn.execute(
        "SELECT * FROM unidades WHERE asignatura_id=? ORDER BY orden", (asig_id,)
    ).fetchall()]
    metodologias = [dict(r) for r in conn.execute(
        "SELECT * FROM metodologias WHERE asignatura_id=?", (asig_id,)
    ).fetchall()]
    evaluaciones = [dict(r) for r in conn.execute(
        "SELECT * FROM evaluaciones WHERE asignatura_id=? ORDER BY id", (asig_id,)
    ).fetchall()]
    conn.row_factory = None
    ras = conn.execute("""
        SELECT ra.codigo_completo, ra.descripcion, c.codigo, c.tipo
        FROM tributaciones ar
        JOIN resultados_aprendizaje ra ON ra.id=ar.ra_id
        JOIN niveles_dominio nd ON nd.id = ra.nivel_dominio_id
        JOIN competencias c ON c.id=nd.competencia_id
        WHERE ar.asignatura_id=?
        ORDER BY c.codigo, ra.codigo_completo
    """, (asig_id,)).fetchall()
    conn.close()
    ras = [{"codigo_completo": r[0], "descripcion": r[1],
            "comp": r[2], "tipo": r[3]} for r in ras]
    return asig, unidades, metodologias, evaluaciones, ras

def buscar_word_original(codigo):
    """Busca el .docx original por código de asignatura."""
    codigo_norm = codigo.strip().replace(" ", "")
    for docx in PROG_DIR.rglob("*.docx"):
        # Normalizar nombre de archivo para comparar
        nombre = docx.stem.replace(" ", "").upper()
        if codigo_norm.upper() in nombre:
            return docx
    return None

# ── Helpers XML ───────────────────────────────────────────────────

def get_cell_text(cell):
    return "".join(t.text for t in cell._tc.iter(f"{{{WNS}}}t") if t.text)

def clear_cell(cell):
    """Vacía el contenido de una celda conservando su formato."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    # Dejar al menos un párrafo vacío
    if not cell.paragraphs:
        cell.add_paragraph()

def set_cell_text(cell, text, bold=False, preserve_style=True):
    """Reemplaza el texto de una celda conservando el estilo del primer run."""
    para = cell.paragraphs[0]
    # Copiar formato del primer run si existe
    fmt = {}
    if para.runs and preserve_style:
        r0 = para.runs[0]
        fmt = {
            "bold":   r0.bold,
            "italic": r0.italic,
            "size":   r0.font.size,
            "color":  r0.font.color.rgb if r0.font.color and r0.font.color.type else None,
            "name":   r0.font.name,
        }
    # Borrar runs existentes
    for run in para.runs:
        run.text = ""
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run()
    run.text = text
    if bold:
        run.bold = True
    elif fmt.get("bold") is not None:
        run.bold = fmt["bold"]

def clone_paragraph_format(src_para, dst_para):
    """Copia el pPr (formato de párrafo) de src a dst."""
    src_pPr = src_para._p.find(qn("w:pPr"))
    if src_pPr is not None:
        dst_pPr = dst_para._p.find(qn("w:pPr"))
        if dst_pPr is not None:
            dst_para._p.remove(dst_pPr)
        dst_para._p.insert(0, copy.deepcopy(src_pPr))

def clone_run_format(src_run, dst_run):
    """Copia el rPr (formato de run) de src a dst."""
    src_rPr = src_run._r.find(qn("w:rPr"))
    if src_rPr is not None:
        dst_rPr = dst_run._r.find(qn("w:rPr"))
        if dst_rPr is not None:
            dst_run._r.remove(dst_rPr)
        dst_run._r.insert(0, copy.deepcopy(src_rPr))

def get_body_elements(doc):
    """Retorna lista de elementos del body con su índice."""
    return list(doc.element.body)

def find_section_index(body_elems, marker_text):
    """Encuentra el índice del primer elemento que contiene marker_text."""
    for i, elem in enumerate(body_elems):
        tag = elem.tag.split("}")[-1]
        if tag == "p":
            text = "".join(t.text for t in elem.iter(f"{{{WNS}}}t") if t.text)
            if marker_text.lower() in text.lower():
                return i
    return -1

def find_table_after(body_elems, start_idx):
    """Encuentra la primera tabla después del índice dado."""
    for i in range(start_idx, len(body_elems)):
        if body_elems[i].tag.split("}")[-1] == "tbl":
            return i
    return -1

# ── Edición de secciones ──────────────────────────────────────────

def editar_tabla_identificacion(doc, asig):
    """Tabla de identificación: nombre, código, requisitos."""
    body = get_body_elements(doc)
    idx_sec = find_section_index(body, "IDENTIFICACIÓN:")
    if idx_sec < 0:
        return
    # La tabla de identificación simple (Nombre / Código / Requisito) viene después
    # Hay dos tablas de identificación: la grande (Facultad/Carrera) y la pequeña
    # Buscar la que tiene "Nombre" y "Código" como primera columna
    tables_after = [i for i in range(idx_sec, len(body))
                    if body[i].tag.split("}")[-1] == "tbl"]

    for tbl_idx in tables_after[:3]:
        tbl_elem = body[tbl_idx]
        rows = list(tbl_elem.iter(f"{{{WNS}}}tr"))
        if len(rows) >= 2:
            first_cell = "".join(t.text for t in
                list(rows[0].iter(f"{{{WNS}}}tc"))[0].iter(f"{{{WNS}}}t") if t.text)
            if "Nombre" in first_cell or "nombre" in first_cell:
                # Esta es la tabla pequeña de identificación
                for row in rows:
                    cells = list(row.iter(f"{{{WNS}}}tc"))
                    if len(cells) >= 2:
                        label = "".join(t.text for t in cells[0].iter(f"{{{WNS}}}t") if t.text)
                        if "Nombre" in label:
                            set_cell_text(
                                _tc_to_cell(doc, cells[1]),
                                asig.get("nombre", "")
                            )
                        elif "Código" in label or "Codigo" in label:
                            set_cell_text(
                                _tc_to_cell(doc, cells[1]),
                                asig.get("codigo", "")
                            )
                        elif "Requisito" in label:
                            set_cell_text(
                                _tc_to_cell(doc, cells[1]),
                                asig.get("requisitos", "") or "Sin Requisito"
                            )
                break

def _tc_to_cell(doc, tc_elem):
    """Wrapper para obtener un Cell desde un elemento tc XML."""
    from docx.table import _Cell
    return _Cell(tc_elem, None)

def editar_ras(doc, ras):
    """
    Reemplaza los párrafos de RAs entre 'RESULTADOS DE APRENDIZAJE Y DESEMPEÑOS'
    y 'UNIDADES DE APRENDIZAJE Y CONTENIDOS'.
    """
    body = get_body_elements(doc)
    idx_ra  = find_section_index(body, "RESULTADOS DE APRENDIZAJE")
    idx_uni = find_section_index(body, "UNIDADES DE APRENDIZAJE")
    if idx_ra < 0 or idx_uni < 0:
        return

    # El texto introductorio está justo después del header (idx_ra+1)
    # Los párrafos de RAs están entre idx_ra+2 y idx_uni-1
    # Identificar el párrafo introductorio y los párrafos de lista de RAs
    ra_parrafos_idx = []
    for i in range(idx_ra + 1, idx_uni):
        elem = body[i]
        if elem.tag.split("}")[-1] == "p":
            text = "".join(t.text for t in elem.iter(f"{{{WNS}}}t") if t.text)
            # Los párrafos de RAs tienen formato "CG1, D1:" o "CL1, N1, RA1:"
            if re.match(r"[A-Z]{2}\d", text.strip()):
                ra_parrafos_idx.append(i)

    if not ra_parrafos_idx:
        return

    # Tomar el formato del primer párrafo de RA como plantilla
    plantilla_elem = body[ra_parrafos_idx[0]]

    # Eliminar todos los párrafos de RA existentes (de atrás para adelante)
    body_xml = doc.element.body
    for i in sorted(ra_parrafos_idx, reverse=True):
        body_xml.remove(body[i])

    # Reinsertar párrafos con los RAs de la BD
    # Insertar antes de idx_uni (que ahora cambió de posición)
    body_refresh = get_body_elements(doc)
    idx_uni_new = find_section_index(body_refresh, "UNIDADES DE APRENDIZAJE")

    for ra in reversed(ras):
        texto = ra["codigo_completo"]
        if ra.get("descripcion"):
            texto += f": {ra['descripcion']}"

        # Clonar párrafo plantilla
        nuevo_p = copy.deepcopy(plantilla_elem)
        # Limpiar runs y poner texto nuevo
        for r in nuevo_p.findall(f".//{{{WNS}}}r"):
            nuevo_p.remove(r)
        run_elem = copy.deepcopy(
            plantilla_elem.findall(f".//{{{WNS}}}r")[0]
            if plantilla_elem.findall(f".//{{{WNS}}}r") else OxmlElement("w:r")
        )
        t_elem = run_elem.find(f"{{{WNS}}}t")
        if t_elem is None:
            t_elem = OxmlElement("w:t")
            run_elem.append(t_elem)
        t_elem.text = texto
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        nuevo_p.append(run_elem)

        ref_elem = body_refresh[idx_uni_new]
        doc.element.body.insert(
            list(doc.element.body).index(ref_elem),
            nuevo_p
        )

def editar_unidades(doc, unidades):
    """Reemplaza filas de datos en la tabla de Unidades y Contenidos."""
    body = get_body_elements(doc)
    idx_uni = find_section_index(body, "UNIDADES DE APRENDIZAJE")
    if idx_uni < 0:
        return
    tbl_idx = find_table_after(body, idx_uni)
    if tbl_idx < 0:
        return

    # Obtener el objeto Table de python-docx
    tabla = _get_table_obj(doc, body[tbl_idx])
    if tabla is None:
        return

    # Guardar fila header (primera fila)
    header_row = tabla.rows[0]

    # Eliminar filas de datos (todo menos el header)
    tbl_elem = body[tbl_idx]
    rows_xml = tbl_elem.findall(f"{{{WNS}}}tr")
    for row_xml in rows_xml[1:]:
        tbl_elem.remove(row_xml)

    # Plantilla de fila (clonar la segunda fila si existía, si no usar la primera)
    plantilla_row = rows_xml[1] if len(rows_xml) > 1 else rows_xml[0]

    for u in unidades:
        nueva_fila = copy.deepcopy(plantilla_row)
        celdas = nueva_fila.findall(f"{{{WNS}}}tc")
        if len(celdas) >= 2:
            # Columna 1: indicador de logro / RA
            _set_tc_text(celdas[0], u.get("indicador_logro", "") or "")
            # Columna 2: contenidos
            _set_tc_text(celdas[1], u.get("contenidos", "") or u.get("nombre", ""))
        tbl_elem.append(nueva_fila)

def editar_evaluaciones(doc, evaluaciones):
    """Reemplaza filas de datos en la tabla de Evaluaciones."""
    body = get_body_elements(doc)
    idx_ev = find_section_index(body, "ESTRATEGIA DE EVALUACIÓN")
    if idx_ev < 0:
        idx_ev = find_section_index(body, "EVALUACIÓN")
    if idx_ev < 0:
        return
    tbl_idx = find_table_after(body, idx_ev)
    if tbl_idx < 0:
        return

    tbl_elem = body[tbl_idx]
    rows_xml = tbl_elem.findall(f"{{{WNS}}}tr")
    if len(rows_xml) < 2:
        return

    plantilla_row = rows_xml[1]

    # Eliminar filas de datos
    for row_xml in rows_xml[1:]:
        tbl_elem.remove(row_xml)

    for ev in evaluaciones:
        tipo = ev.get("tipo", "")
        porc = ev.get("porcentaje", "")
        if not tipo.strip():
            continue
        nueva_fila = copy.deepcopy(plantilla_row)
        celdas = nueva_fila.findall(f"{{{WNS}}}tc")
        if len(celdas) >= 2:
            _set_tc_text(celdas[0], tipo)
            _set_tc_text(celdas[1], porc)
        tbl_elem.append(nueva_fila)

def editar_datos_actualizacion(doc, asig):
    """Actualiza la tabla de Datos de Actualización."""
    body = get_body_elements(doc)
    idx = find_section_index(body, "DATOS ACTUALIZACIÓN")
    if idx < 0:
        return
    tbl_idx = find_table_after(body, idx)
    if tbl_idx < 0:
        return

    tbl_elem = body[tbl_idx]
    for row in tbl_elem.findall(f"{{{WNS}}}tr"):
        celdas = row.findall(f"{{{WNS}}}tc")
        if len(celdas) >= 2:
            label = _get_tc_text(celdas[0])
            if "Versión" in label or "Fecha" in label:
                version = asig.get("version", "") or f"V1.0 {datetime.now().year}"
                _set_tc_text(celdas[1], version)

# ── Helpers XML internos ──────────────────────────────────────────

def _get_tc_text(tc_elem):
    return "".join(t.text for t in tc_elem.iter(f"{{{WNS}}}t") if t.text)

def _set_tc_text(tc_elem, text):
    """Pone texto en un elemento tc XML, conservando formato del primer run."""
    paras = tc_elem.findall(f"{{{WNS}}}p")
    if not paras:
        return
    para = paras[0]
    runs = para.findall(f"{{{WNS}}}r")

    # Limpiar todos los párrafos extra
    for p in paras[1:]:
        tc_elem.remove(p)

    # Limpiar runs del primer párrafo
    for r in runs:
        para.remove(r)

    # Crear run nuevo con el texto
    run = OxmlElement("w:r")
    # Copiar rPr del primer run original si existía
    if runs:
        rPr = runs[0].find(f"{{{WNS}}}rPr")
        if rPr is not None:
            run.append(copy.deepcopy(rPr))
    t = OxmlElement("w:t")
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)
    para.append(run)

def _get_table_obj(doc, tbl_elem):
    """Obtiene el objeto Table de python-docx a partir del elemento XML."""
    for t in doc.tables:
        if t._tbl is tbl_elem:
            return t
    return None

# ── Función principal ─────────────────────────────────────────────

def _get_programa_data(asignatura_id):
    """Obtiene todos los datos de una asignatura desde la BD."""
    conn = sqlite3.connect(str(RUTA_DB))
    conn.row_factory = sqlite3.Row

    asig = conn.execute("SELECT * FROM asignaturas WHERE id=?", (asignatura_id,)).fetchone()
    if not asig:
        conn.close()
        return None

    responsables = conn.execute(
        "SELECT rol, nombre FROM responsables WHERE asignatura_id=? ORDER BY rol",
        (asignatura_id,)
    ).fetchall()

    ras = conn.execute("""
        SELECT ra.codigo_completo, ra.descripcion, c.tipo
        FROM tributaciones ar
        JOIN resultados_aprendizaje ra ON ra.id = ar.ra_id
        JOIN niveles_dominio nd ON nd.id = ra.nivel_dominio_id
        JOIN competencias c ON c.id = nd.competencia_id
        WHERE ar.asignatura_id = ?
        ORDER BY c.codigo, nd.codigo_nivel, ra.codigo_ra
    """, (asignatura_id,)).fetchall()

    # Competencias+NDs con tributaciones (para la sección de aporte visual).
    # comp_desc usa la descripción del ND más avanzado disponible como definición
    # canónica de la competencia (ya que competencias.descripcion está vacío).
    aporte_rows = conn.execute("""
        SELECT DISTINCT c.codigo,
               (SELECT nd2.descripcion FROM niveles_dominio nd2
                WHERE nd2.competencia_id = c.id
                ORDER BY nd2.codigo_nivel DESC LIMIT 1) AS comp_desc,
               c.tipo,
               nd.codigo_nivel, nd.descripcion AS nd_desc
        FROM tributaciones ar
        JOIN resultados_aprendizaje ra ON ra.id = ar.ra_id
        JOIN niveles_dominio nd ON nd.id = ra.nivel_dominio_id
        JOIN competencias c ON c.id = nd.competencia_id
        WHERE ar.asignatura_id = ?
        ORDER BY c.tipo, c.codigo, nd.codigo_nivel
    """, (asignatura_id,)).fetchall()

    unidades = conn.execute(
        "SELECT orden, nombre, contenidos, indicador_logro FROM unidades WHERE asignatura_id=? ORDER BY orden",
        (asignatura_id,)
    ).fetchall()

    metodologias = conn.execute(
        "SELECT descripcion FROM metodologias WHERE asignatura_id=?",
        (asignatura_id,)
    ).fetchall()

    evaluaciones = conn.execute(
        "SELECT tipo, porcentaje FROM evaluaciones WHERE asignatura_id=?",
        (asignatura_id,)
    ).fetchall()

    bibliografia_basica = conn.execute(
        "SELECT numero, autor, titulo, editorial, anio, isbn, ejemplares FROM bibliografia WHERE asignatura_id=? AND tipo='basica' ORDER BY id",
        (asignatura_id,)
    ).fetchall()

    bibliografia_comp = conn.execute(
        "SELECT numero, autor, titulo, editorial, anio, isbn, ejemplares FROM bibliografia WHERE asignatura_id=? AND tipo='complementaria' ORDER BY id",
        (asignatura_id,)
    ).fetchall()

    linkografia = conn.execute(
        """SELECT tipo_documento, autor, titulo_articulo, anio,
                  titulo_revista, volumen, url, disponible_en
           FROM linkografia WHERE asignatura_id=? ORDER BY id""",
        (asignatura_id,)
    ).fetchall()

    conn.close()

    # Agrupar aporte_rows en bloques por tipo → competencia → niveles
    _TIPO_ORDER = {"licenciatura": 0, "titulo": 1, "sello_uv": 2}
    _TIPO_INTRO = {
        "licenciatura": (
            "Esta asignatura aporta al perfil de Licenciatura a través de resultados de "
            "aprendizaje que tributan a las siguientes competencias específicas de licenciatura "
            "en el primer nivel de dominio:"
        ),
        "titulo": (
            "Esta asignatura aporta al perfil de Título a través de resultados de aprendizaje "
            "que tributan a las siguientes competencias específicas del título profesional:"
        ),
        "sello_uv": (
            "Esta asignatura aporta al logro del perfil de egreso y licenciatura a través de "
            "indicadores de desempeño que tributan a las siguientes competencias genéricas en su "
            "nivel inicial de desempeño:"
        ),
    }
    _TIPO_LABEL = {
        "licenciatura": "Competencia de Licenciatura:",
        "titulo":       "Competencia de Título:",
        "sello_uv":     "Competencia Genérica:",
    }
    _bloques: dict = {}
    for row in aporte_rows:
        comp_cod, comp_desc, tipo, nd_cod, nd_desc = row
        if tipo not in _bloques:
            _bloques[tipo] = {}
        if comp_cod not in _bloques[tipo]:
            _bloques[tipo][comp_cod] = {"descripcion": comp_desc or "", "niveles": []}
        _bloques[tipo][comp_cod]["niveles"].append({"codigo": nd_cod, "descripcion": nd_desc or ""})

    aporte_bloques = []
    for tipo in sorted(_bloques, key=lambda t: _TIPO_ORDER.get(t, 99)):
        comps = []
        for comp_cod, comp_data in _bloques[tipo].items():
            comps.append({"codigo": comp_cod, "descripcion": comp_data["descripcion"],
                          "niveles": comp_data["niveles"]})
        aporte_bloques.append({
            "tipo": tipo,
            "intro": _TIPO_INTRO.get(tipo, ""),
            "label": _TIPO_LABEL.get(tipo, "Competencia:"),
            "competencias": comps,
        })

    return {
        "asig": dict(asig),
        "responsables": {r["rol"]: r["nombre"] for r in responsables},
        "ras": [dict(r) for r in ras],
        "aporte_bloques": aporte_bloques,
        "unidades": [dict(u) for u in unidades],
        "metodologias": [m["descripcion"] for m in metodologias],
        "evaluaciones": [dict(e) for e in evaluaciones],
        "bibliografia_basica": [dict(b) for b in bibliografia_basica],
        "bibliografia_comp": [dict(b) for b in bibliografia_comp],
        "linkografia": [dict(l) for l in linkografia],
    }


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_para(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text or "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def _heading(doc, text, level=1):
    """Encabezado UV (fondo azul, texto blanco)."""
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "1F4E79")
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "Calibri"
    doc.add_paragraph()


def _section_title(doc, text):
    """Título de sección (bold, azul)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = UV_BLUE


def _add_cell_para(cell, indent_cm=0.0):
    """Añade un párrafo nuevo al final de una celda y lo retorna."""
    p = OxmlElement("w:p")
    if indent_cm:
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(int(indent_cm * 567)))  # 1 cm ≈ 567 twips
        pPr.append(ind)
        p.append(pPr)
    cell._tc.append(p)
    from docx.text.paragraph import Paragraph as _Paragraph
    return _Paragraph(p, cell)


def _run(para, text, bold=False, underline=False, size=10, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return run


def _boxed_text(doc, text, size=10):
    """Envuelve un bloque de texto plano en una celda con borde (Table Grid 1×1)."""
    if not (text or "").strip():
        return
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run(text.strip())
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def _boxed_ra_list(doc, intro_text, ras, tipo_color_map):
    """
    Envuelve la sección de RAs (intro + lista) en una celda con borde.
    ras: lista de dicts con 'codigo_completo', 'descripcion', 'tipo'
    """
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)

    # Intro
    p0 = cell.paragraphs[0]
    _run(p0, intro_text, size=10)

    # Bullets
    for ra in ras:
        p_ra = _add_cell_para(cell)
        _run(p_ra, "• ", size=10)
        r_cod = p_ra.add_run(ra["codigo_completo"])
        r_cod.bold = True
        r_cod.font.size = Pt(10)
        r_cod.font.name = "Calibri"
        tipo = ra.get("tipo", "licenciatura")
        r_cod.font.color.rgb = tipo_color_map.get(tipo, UV_BLUE)
        if ra.get("descripcion"):
            _run(p_ra, f": {ra['descripcion']}", size=10)


def _render_aporte_perfil(doc, bloques):
    """
    Renderiza la sección 'APORTE AL PERFIL DE EGRESO' como una tabla de
    1 columna con bordes negros, donde cada fila es un bloque por tipo de
    competencia (licenciatura / título / sello UV), separados por línea
    horizontal interna, imitando el diseño de la plantilla UV.
    """
    if not bloques:
        p = doc.add_paragraph()
        _run(p, "Sin tributaciones registradas.", size=10)
        return

    tbl = doc.add_table(rows=len(bloques), cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, bloque in enumerate(bloques):
        cell = tbl.cell(i, 0)

        # ── Párrafo 1: checkmark + texto introductorio ─────────────
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after  = Pt(2)
        _run(p0, "✓  ", bold=True, size=10)   # ✓
        _run(p0, bloque["intro"], size=10)

        # ── Párrafo vacío de separación ────────────────────────────
        _add_cell_para(cell)

        # ── Etiqueta de competencia (subrayada) ─────────────────────
        p_label = _add_cell_para(cell)
        _run(p_label, bloque["label"], underline=True, size=10)

        # ── Cada competencia ────────────────────────────────────────
        for comp in bloque["competencias"]:
            # "- CL1. descripcion"
            p_comp = _add_cell_para(cell, indent_cm=0.5)
            _run(p_comp, "- ", size=10)
            _run(p_comp, f"{comp['codigo']}.", bold=True, size=10)
            if comp["descripcion"]:
                _run(p_comp, f" {comp['descripcion']}", size=10)

            # "  - N1. descripcion"
            for nd in comp["niveles"]:
                nd_label = nd["codigo"].replace("ND", "N")
                p_nd = _add_cell_para(cell, indent_cm=1.25)
                _run(p_nd, "- ", size=10)
                _run(p_nd, f"{nd_label}.", bold=True, size=10)
                if nd["descripcion"]:
                    _run(p_nd, f" {nd['descripcion']}", size=10)

        # Pequeño margen al final de la celda
        _add_cell_para(cell)


def generar_programa_individual(asignatura_id, salida=None):
    """Genera el documento Word del programa de asignatura desde la BD."""
    data = _get_programa_data(asignatura_id)
    if not data:
        raise ValueError(f"Asignatura {asignatura_id} no encontrada")

    asig = data["asig"]
    responsables = data["responsables"]
    codigo = (asig["codigo"] or "asig").replace(" ", "_")

    RUTA_OUTPUT.mkdir(parents=True, exist_ok=True)
    if salida is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        salida = str(RUTA_OUTPUT / f"programa_{codigo}_{ts}.docx")
    ruta = salida

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── TÍTULO PRINCIPAL ──────────────────────────────────────────
    t_titulo = doc.add_table(rows=1, cols=2)
    t_titulo.style = "Table Grid"
    c_logo = t_titulo.cell(0, 0)
    c_logo.width = Cm(5)
    _set_cell_bg(c_logo, "FFFFFF")
    if LOGO_PATH.exists():
        p = c_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(4.5))
    else:
        _cell_para(c_logo, "UV", bold=True, size=16, color=RGBColor(0x1F, 0x4E, 0x79), align=WD_ALIGN_PARAGRAPH.CENTER)
    c_title = t_titulo.cell(0, 1)
    _set_cell_bg(c_title, "1F4E79")
    _cell_para(c_title, "PROGRAMA DE ASIGNATURA UV", bold=True, size=14,
               color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ── SECCIÓN 1: DESCRIPCIÓN GENERAL ───────────────────────────
    _heading(doc, "DESCRIPCIÓN GENERAL DE LA ASIGNATURA")

    _section_title(doc, "IDENTIFICACIÓN DE LA ASIGNATURA:")

    t_id = doc.add_table(rows=7, cols=6)
    t_id.style = "Table Grid"

    def _cell(r, c):
        return t_id.cell(r, c)

    _cell_para(_cell(0, 0), "Facultad:", bold=True, size=9)
    _cell_para(_cell(0, 1), asig.get("facultad", "") or "", size=9)
    _cell(0, 1).merge(_cell(0, 2))
    _cell_para(_cell(0, 3), "Carrera:", bold=True, size=9)
    _cell_para(_cell(0, 4), asig.get("carrera", "") or "", size=9)
    _cell(0, 4).merge(_cell(0, 5))

    _cell_para(_cell(1, 0), "Nombre:", bold=True, size=9)
    _cell_para(_cell(1, 1), asig.get("nombre", "") or "", size=9)
    _cell(1, 1).merge(_cell(1, 2))
    _cell_para(_cell(1, 3), "Código:", bold=True, size=9)
    _cell_para(_cell(1, 4), asig.get("codigo", "") or "", size=9)
    _cell(1, 4).merge(_cell(1, 5))

    _cell_para(_cell(2, 0), "Nivel:", bold=True, size=9)
    _cell_para(_cell(2, 1), asig.get("nivel", "") or "", size=9)
    _cell(2, 1).merge(_cell(2, 2))
    _cell_para(_cell(2, 3), "Duración:", bold=True, size=9)
    _cell_para(_cell(2, 4), asig.get("duracion", "") or "", size=9)
    _cell(2, 4).merge(_cell(2, 5))

    _cell_para(_cell(3, 0), "Requisito(s):", bold=True, size=9)
    req_cell = _cell(3, 1)
    req_cell.merge(_cell(3, 5))
    _cell_para(req_cell, asig.get("requisitos", "") or "", size=9)

    headers_h = ["Horas cronológicas semanales", "", "", "N° de semanas", "Total horas semestrales", "N° de créditos"]
    for c, h in enumerate(headers_h):
        _cell_para(_cell(4, c), h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(4, 0).merge(_cell(4, 2))

    sub_h = ["Docencia Directa\n(A)", "Trabajo Autónomo\n(B)", "Total\n(C=A+B)", "N° de semanas\n(D)", "Total horas\n(E=C*D)", "N° créditos\n(F=E/27)"]
    for c, h in enumerate(sub_h):
        _cell_para(_cell(5, c), h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    hd = asig.get("horas_directa")
    ha = asig.get("horas_autonoma")
    sem = asig.get("semanas")
    cred = asig.get("creditos")
    total_h = (hd or 0) + (ha or 0)
    # Usar total_horas leído directamente del documento (campo E) cuando está
    # disponible, ya que algunos documentos tienen A y B incorrectos (ej. IMAT 612).
    total_sem = asig.get("total_horas") or (total_h * sem if (total_h and sem) else None)
    vals = [
        str(hd).replace(".", ",") if hd else "",
        str(ha).replace(".", ",") if ha else "",
        str(total_h).replace(".", ",") if total_h else "",
        str(sem) if sem else "",
        str(int(total_sem)) if total_sem else "",
        str(cred) if cred else "",
    ]
    for c, v in enumerate(vals):
        _cell_para(_cell(6, c), v, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    if asig.get("descripcion"):
        _section_title(doc, "DESCRIPCIÓN DE LA ASIGNATURA:")
        _boxed_text(doc, asig["descripcion"])
        doc.add_paragraph()

    _section_title(doc, "APORTE AL PERFIL DE EGRESO:")
    _render_aporte_perfil(doc, data.get("aporte_bloques", []))
    doc.add_paragraph()

    # ── SECCIÓN 2: PROGRAMA DE LA ASIGNATURA ─────────────────────
    _heading(doc, "PROGRAMA DE LA ASIGNATURA")

    _section_title(doc, "IDENTIFICACIÓN DE LA ASIGNATURA:")
    t_id2 = doc.add_table(rows=3, cols=2)
    t_id2.style = "Table Grid"
    rows2 = [("Nombre", asig.get("nombre", "")), ("Código", asig.get("codigo", "")), ("Requisito(s)", asig.get("requisitos", ""))]
    for i, (k, v) in enumerate(rows2):
        _cell_para(t_id2.cell(i, 0), k, bold=True, size=9)
        _cell_para(t_id2.cell(i, 1), v or "", size=9)
    doc.add_paragraph()

    _section_title(doc, "RESULTADOS DE APRENDIZAJE Y DESEMPEÑOS:")
    _boxed_ra_list(
        doc,
        "Al final de la asignatura los estudiantes serán capaces de demostrar los siguientes resultados:",
        data["ras"],
        TIPO_COLOR,
    )
    doc.add_paragraph()

    _section_title(doc, "UNIDADES DE APRENDIZAJE Y CONTENIDOS:")
    if data["unidades"]:
        t_uni = doc.add_table(rows=1, cols=2)
        t_uni.style = "Table Grid"
        for ci, h in enumerate(["Resultado de aprendizaje y/o Desempeños",
                                 "Unidades de Aprendizaje y contenidos"]):
            c = t_uni.cell(0, ci)
            _set_cell_bg(c, "1F4E79")
            _cell_para(c, h, bold=True, size=9, color=RGBColor(255, 255, 255),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for u in data["unidades"]:
            row = t_uni.add_row()
            _cell_para(row.cells[0], u.get("indicador_logro", "") or "", size=9)
            _cell_para(row.cells[1], (u.get("contenidos", "") or "").strip(), size=9)
    doc.add_paragraph()

    # Sección "3.1 EXPERIENCIAS DE LABORATORIO" si existe
    lab_text = (data["asig"].get("experiencias_laboratorio") or "").strip()
    if lab_text:
        _section_title(doc, "3.1 EXPERIENCIAS DE LABORATORIO:")
        _boxed_text(doc, lab_text)
        doc.add_paragraph()

    _section_title(doc, "METODOLOGÍA O ESTRATEGIA DE ENSEÑANZA - APRENDIZAJE:")
    # Lista estándar UV de metodologías (formato tabla de checkboxes 4 columnas)
    _UV_METODS = [
        ("Procedimiento de Pausas.",                  "Aprendizaje por descubrimiento."),
        ("Estudio de Casos.",                          "Aprendizaje basado en equipos."),
        ("Juego de Roles.",                            "Clase expositiva activa."),
        ("Aprendizaje colaborativo o cooperativo.",    "Simulación."),
        ("Aprendizaje basado en Problemas (ABP).",     "Tutorías."),
        ("Aprendizaje basado en Proyectos.",           "Salidas a terreno."),
        ("Aprendizaje Servicio.",                      "Otro, especifique:"),
        ("Aprendizaje Invertido.",                     ""),
    ]
    # Split each methodology block by newlines to get individual items
    metod_items = []
    for bloque in data["metodologias"]:
        for line in bloque.split('\n'):
            line = line.strip()
            if line:
                metod_items.append(line)

    seleccionadas = {m.rstrip('.').lower() for m in metod_items}

    def _esta_marcada(nombre):
        n = nombre.strip().rstrip('.').lower()
        return any(n == s or (len(n) > 8 and (n in s or s in n)) for s in seleccionadas)

    # Only use UV checkbox table when the majority of items match the template
    n_uv_total = sum(1 for izq, der in _UV_METODS if izq or der)
    n_marcadas = sum(1 for izq, der in _UV_METODS
                     if _esta_marcada(izq) or (der and _esta_marcada(der)))

    if n_marcadas >= 3:
        # Render UV checkbox table
        t_met = doc.add_table(rows=len(_UV_METODS), cols=4)
        t_met.style = "Table Grid"
        for ri, (izq, der) in enumerate(_UV_METODS):
            row = t_met.rows[ri]
            _cell_para(row.cells[0], izq, size=9)
            _cell_para(row.cells[1], "X" if _esta_marcada(izq) else "", size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_para(row.cells[2], der, size=9)
            _cell_para(row.cells[3], "X" if (der and _esta_marcada(der)) else "", size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
    elif metod_items:
        # Free-text methodology: enclosed box (matches official document format)
        _boxed_text(doc, "\n".join(metod_items))
    doc.add_paragraph()

    _section_title(doc, "METODOLOGÍA O ESTRATEGIA DE EVALUACIÓN:")
    if data["evaluaciones"]:
        t_ev = doc.add_table(rows=1, cols=2)
        t_ev.style = "Table Grid"
        for ci, h in enumerate(["Tipo de evaluación:", "Porcentaje (%) que corresponde:"]):
            c = t_ev.cell(0, ci)
            _set_cell_bg(c, "1F4E79")
            _cell_para(c, h, bold=True, size=9, color=RGBColor(255, 255, 255))
        for ev in data["evaluaciones"]:
            row = t_ev.add_row()
            _cell_para(row.cells[0], ev.get("tipo", "") or "", size=9)
            _cell_para(row.cells[1], ev.get("porcentaje", "") or "", size=9)
    # Texto libre después de la tabla de evaluaciones (política de notas, IA, etc.)
    desc_ev = (data["asig"].get("descripcion_evaluaciones") or "").strip()
    if desc_ev:
        if data["evaluaciones"]:
            # Añade como fila fusionada al final de la tabla de evaluaciones
            row_desc = t_ev.add_row()
            merged = row_desc.cells[0].merge(row_desc.cells[1])
            p_desc = merged.paragraphs[0]
            run_desc = p_desc.add_run(desc_ev)
            run_desc.font.size = Pt(9)
            run_desc.font.name = "Calibri"
        else:
            _boxed_text(doc, desc_ev)
    doc.add_paragraph()

    _BIBLIO_HDRS = ["Autor", "Título", "Editorial", "Año", "ISBN", "Nº Ejemplares disponibles"]
    _BIBLIO_KEYS = ["autor", "titulo", "editorial", "anio", "isbn", "ejemplares"]

    if data["bibliografia_basica"]:
        _section_title(doc, "BIBLIOGRAFÍA BÁSICA OBLIGATORIA:")
        t_bib = doc.add_table(rows=1, cols=6)
        t_bib.style = "Table Grid"
        for ci, h in enumerate(_BIBLIO_HDRS):
            c = t_bib.cell(0, ci)
            _set_cell_bg(c, "1F4E79")
            _cell_para(c, h, bold=True, size=8, color=RGBColor(255, 255, 255))
        for b in data["bibliografia_basica"]:
            row = t_bib.add_row()
            for ci, k in enumerate(_BIBLIO_KEYS):
                _cell_para(row.cells[ci], b.get(k, "") or "", size=8)
        doc.add_paragraph()

    if data["bibliografia_comp"]:
        _section_title(doc, "BIBLIOGRAFÍA COMPLEMENTARIA:")
        t_bib2 = doc.add_table(rows=1, cols=6)
        t_bib2.style = "Table Grid"
        for ci, h in enumerate(_BIBLIO_HDRS):
            c = t_bib2.cell(0, ci)
            _set_cell_bg(c, "1F4E79")
            _cell_para(c, h, bold=True, size=8, color=RGBColor(255, 255, 255))
        for b in data["bibliografia_comp"]:
            row = t_bib2.add_row()
            for ci, k in enumerate(_BIBLIO_KEYS):
                _cell_para(row.cells[ci], b.get(k, "") or "", size=8)
        doc.add_paragraph()

    if data["linkografia"]:
        _section_title(doc, "LINKOGRAFÍA:")
        _LINK_HDRS = ["Tipo de Documento", "Autor", "Título Artículo / Documento / Sitio Web",
                      "Año", "Título e-Revista / e-Libro", "Vol(Nº)",
                      "Dirección Electrónica (URL)", "Disponible en"]
        _LINK_KEYS = ["tipo_documento", "autor", "titulo_articulo", "anio",
                      "titulo_revista", "volumen", "url", "disponible_en"]
        t_lnk = doc.add_table(rows=1, cols=8)
        t_lnk.style = "Table Grid"
        for ci, h in enumerate(_LINK_HDRS):
            c = t_lnk.cell(0, ci)
            _set_cell_bg(c, "1F4E79")
            _cell_para(c, h, bold=True, size=8, color=RGBColor(255, 255, 255),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for lk in data["linkografia"]:
            row = t_lnk.add_row()
            for ci, k in enumerate(_LINK_KEYS):
                _cell_para(row.cells[ci], lk.get(k, "") or "", size=8)
        doc.add_paragraph()

    if asig.get("otros_recursos"):
        _section_title(doc, "OTROS RECURSOS:")
        _boxed_text(doc, asig["otros_recursos"])
        doc.add_paragraph()

    _section_title(doc, "DATOS ACTUALIZACIÓN:")
    t_act = doc.add_table(rows=3, cols=2)
    t_act.style = "Table Grid"
    resp = responsables.get("responsable", "") or ""
    doc_cargo = responsables.get("docente_a_cargo", "") or ""
    version = asig.get("version", "") or ""
    for i, (k, v) in enumerate([("Responsable(s) del programa:", resp), ("Docente(s) a cargo:", doc_cargo), ("Versión / Fecha de Actualización:", version)]):
        _cell_para(t_act.cell(i, 0), k, bold=True, size=9)
        _cell_para(t_act.cell(i, 1), v, size=9)

    doc.save(str(ruta))
    print(f"✓ Programa generado: {ruta}")
    return str(ruta)


# ── Mapa de Progreso ──────────────────────────────────────────────

def generar_mapa_progreso(salida=None):
    """Genera el Mapa de Progreso en Word desde la BD."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    LOGO_PATH = Path("data/uv_logo_nuevo.png")

    conn = sqlite3.connect(str(DB_PATH)); conn.row_factory = None

    competencias = conn.execute("""
        SELECT id, codigo, tipo, descripcion FROM competencias
        WHERE tipo != 'desconocido'
        ORDER BY CASE tipo WHEN 'licenciatura' THEN 0 WHEN 'titulo' THEN 1 WHEN 'sello_uv' THEN 2 END, codigo
    """).fetchall()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

    def _para(cell, text, bold=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = align
        run = p.add_run(text); run.bold = bold
        run.font.size = Pt(size)
        if color: run.font.color.rgb = color

    # Header
    t_hdr = doc.add_table(rows=1, cols=2); t_hdr.style = 'Table Grid'
    c_logo = t_hdr.cell(0,0); c_logo.width = Cm(5)
    if LOGO_PATH.exists():
        p = c_logo.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Cm(4.5))
    c_tit = t_hdr.cell(0,1)
    _set_cell_bg(c_tit, "1F4E79")
    _para(c_tit, "MAPA DE PROGRESO\nINGENIERÍA CIVIL MATEMÁTICA — PLAN 2025",
          bold=True, color=RGBColor(255,255,255), size=13,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    TIPO_HDR = {"licenciatura":"COMPETENCIAS DE LICENCIATURA",
                "titulo":"COMPETENCIAS ESPECÍFICAS DEL TÍTULO PROFESIONAL",
                "sello_uv":"COMPETENCIAS GENÉRICAS SELLO UV"}
    TIPO_HEX = {"licenciatura":"1F4E79","titulo":"375623","sello_uv":"843C0C"}

    tipo_actual = None
    for comp_id, comp_cod, comp_tipo, comp_desc in competencias:
        if comp_tipo != tipo_actual:
            tipo_actual = comp_tipo
            p = doc.add_paragraph()
            r = p.add_run(TIPO_HDR.get(tipo_actual, tipo_actual.upper()))
            r.bold = True; r.font.size = Pt(13)
            hx = TIPO_HEX.get(tipo_actual,"1F4E79")
            r.font.color.rgb = RGBColor(int(hx[:2],16),int(hx[2:4],16),int(hx[4:],16))

        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"{comp_cod}: {comp_desc or ''}")
        r2.bold = True; r2.font.size = Pt(11)
        hx = TIPO_HEX.get(comp_tipo,"1F4E79")
        r2.font.color.rgb = RGBColor(int(hx[:2],16),int(hx[2:4],16),int(hx[4:],16))

        ras_raw = conn.execute("""
            SELECT nd.codigo_nivel, ra.codigo_completo, ra.descripcion
            FROM resultados_aprendizaje ra
            JOIN niveles_dominio nd ON nd.id = ra.nivel_dominio_id
            WHERE nd.competencia_id=?
            ORDER BY nd.codigo_nivel, ra.codigo_ra
        """, (comp_id,)).fetchall()

        niveles = {}
        for nivel, ccomp, desc in ras_raw:
            niveles.setdefault(nivel or "Sin nivel", []).append((ccomp, desc or ""))

        if not niveles:
            doc.add_paragraph("   (Sin resultados de aprendizaje declarados)").runs[0].italic = True
            doc.add_paragraph(); continue

        cols_niv = sorted(niveles.keys())
        t = doc.add_table(rows=0, cols=len(cols_niv))
        t.style = 'Table Grid'

        # Header
        row_h = t.add_row()
        hx = TIPO_HEX.get(comp_tipo,"1F4E79")
        for ci, niv in enumerate(cols_niv):
            c = row_h.cells[ci]; _set_cell_bg(c, hx)
            _para(c, niv, bold=True, color=RGBColor(255,255,255),
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        # RAs
        max_n = max(len(v) for v in niveles.values())
        for ri in range(max_n):
            row_d = t.add_row()
            for ci, niv in enumerate(cols_niv):
                items = niveles[niv]
                if ri < len(items):
                    ccomp, desc = items[ri]
                    txt = f"• {ccomp}"
                    if desc: txt += f": {desc}"
                    _para(row_d.cells[ci], txt, size=9)

        # Asignaturas por nivel
        row_a = t.add_row()
        bg_claro = {"licenciatura":"DEEAF1","titulo":"E2EFDA","sello_uv":"FCE4D6"}.get(comp_tipo,"F2F2F2")
        for ci, niv in enumerate(cols_niv):
            ra_codes = [r[0] for r in niveles[niv]]
            if ra_codes:
                placeholders = ",".join("?" * len(ra_codes))
                asigs_niv = conn.execute(f"""
                    SELECT DISTINCT a.codigo, a.nombre FROM asignaturas a
                    JOIN tributaciones ar ON ar.asignatura_id=a.id
                    JOIN resultados_aprendizaje ra ON ra.id=ar.ra_id
                    WHERE ra.codigo_completo IN ({placeholders})
                    ORDER BY a.semestre, a.codigo
                """, ra_codes).fetchall()
            else:
                asigs_niv = []
            c = row_a.cells[ci]; _set_cell_bg(c, bg_claro)
            txt = "\n".join(f"• {a[0]} {a[1]}" for a in asigs_niv) or "(ninguna)"
            _para(c, txt, size=8)

        doc.add_paragraph()

    conn.close()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if salida is None:
        salida = str(OUT_DIR / f"mapa_progreso_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
    doc.save(salida)
    print(f"✓ Mapa de Progreso generado: {salida}")
    return salida


def generar_mapa_ra(salida=None):
    """Genera el Mapa de R.A. en Word desde la BD.

    Misma estructura que el Mapa de Progreso, pero cada competencia
    tiene una tabla de 2 columnas: 'Resultado de Aprendizaje' y
    'Asignaturas', con una fila por cada RA.
    """
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    LOGO_PATH = Path("data/uv_logo_nuevo.png")

    conn = sqlite3.connect(str(DB_PATH)); conn.row_factory = None

    competencias = conn.execute("""
        SELECT id, codigo, tipo, descripcion FROM competencias
        WHERE tipo != 'desconocido'
        ORDER BY CASE tipo WHEN 'licenciatura' THEN 0 WHEN 'titulo' THEN 1 WHEN 'sello_uv' THEN 2 END, codigo
    """).fetchall()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

    def _para(cell, text, bold=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        p = cell.paragraphs[0]; p.alignment = align
        run = p.add_run(text); run.bold = bold; run.italic = italic
        run.font.size = Pt(size)
        if color: run.font.color.rgb = color

    def _hex_rgb(hx):
        return RGBColor(int(hx[:2],16), int(hx[2:4],16), int(hx[4:],16))

    # ── Header ────────────────────────────────────────────────────
    t_hdr = doc.add_table(rows=1, cols=2); t_hdr.style = 'Table Grid'
    c_logo = t_hdr.cell(0, 0); c_logo.width = Cm(5)
    if LOGO_PATH.exists():
        p = c_logo.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Cm(4.5))
    c_tit = t_hdr.cell(0, 1)
    _set_cell_bg(c_tit, "1F4E79")
    _para(c_tit, "MAPA DE R.A.\nINGENIERÍA CIVIL MATEMÁTICA — PLAN 2025",
          bold=True, color=RGBColor(255, 255, 255), size=13,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    TIPO_HDR = {
        "licenciatura": "COMPETENCIAS DE LICENCIATURA",
        "titulo":       "COMPETENCIAS ESPECÍFICAS DEL TÍTULO PROFESIONAL",
        "sello_uv":     "COMPETENCIAS GENÉRICAS SELLO UV",
    }
    TIPO_HEX = {"licenciatura": "1F4E79", "titulo": "375623", "sello_uv": "843C0C"}
    BG_CLARO  = {"licenciatura": "DEEAF1", "titulo": "E2EFDA", "sello_uv": "FCE4D6"}

    tipo_actual = None
    for comp_id, comp_cod, comp_tipo, comp_desc in competencias:
        hx     = TIPO_HEX.get(comp_tipo, "1F4E79")
        hx_rgb = _hex_rgb(hx)
        claro  = BG_CLARO.get(comp_tipo, "F2F2F2")

        # Encabezado de tipo (solo cuando cambia)
        if comp_tipo != tipo_actual:
            tipo_actual = comp_tipo
            p = doc.add_paragraph()
            r = p.add_run(TIPO_HDR.get(tipo_actual, tipo_actual.upper()))
            r.bold = True; r.font.size = Pt(13)
            r.font.color.rgb = hx_rgb
            p.paragraph_format.space_before = Pt(6)

        # Línea de competencia: "CL1: descripción"
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"{comp_cod}: {comp_desc or ''}")
        r2.bold = True; r2.font.size = Pt(11)
        r2.font.color.rgb = hx_rgb

        # Todos los RAs de esta competencia, ordenados por ND y código RA
        ras_raw = conn.execute("""
            SELECT nd.codigo_nivel, ra.codigo_completo, ra.descripcion, ra.id
            FROM resultados_aprendizaje ra
            JOIN niveles_dominio nd ON nd.id = ra.nivel_dominio_id
            WHERE nd.competencia_id = ?
            ORDER BY nd.codigo_nivel, ra.codigo_ra
        """, (comp_id,)).fetchall()

        if not ras_raw:
            doc.add_paragraph("   (Sin resultados de aprendizaje declarados)").runs[0].italic = True
            doc.add_paragraph(); continue

        # Tabla 2 columnas: RA | Asignaturas
        t = doc.add_table(rows=0, cols=2)
        t.style = 'Table Grid'

        # Anchura de columnas: RA más ancha
        for cell in t.columns[0].cells if t.rows else []:
            cell.width = Cm(9)

        # Fila de encabezado de columnas
        row_h = t.add_row()
        for ci, lbl in enumerate(["Resultado de Aprendizaje", "Asignaturas"]):
            c = row_h.cells[ci]; _set_cell_bg(c, hx)
            _para(c, lbl, bold=True, size=10,
                  color=RGBColor(255, 255, 255),
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        nivel_actual = None
        for nivel, ccomp, desc, ra_id in ras_raw:
            # Fila separadora por Nivel de Dominio
            if nivel != nivel_actual:
                nivel_actual = nivel
                row_nd = t.add_row()
                c_nd = row_nd.cells[0].merge(row_nd.cells[1])
                _set_cell_bg(c_nd, claro)
                _para(c_nd, f"Nivel de Dominio {nivel}:",
                      bold=True, size=10, color=hx_rgb)

            # Fila con el RA y sus asignaturas
            row_d = t.add_row()

            # Columna RA
            c_ra = row_d.cells[0]
            txt_ra = f"• {ccomp}"
            if desc: txt_ra += f": {desc}"
            _para(c_ra, txt_ra, size=9)

            # Columna Asignaturas
            c_asig = row_d.cells[1]
            asigs = conn.execute("""
                SELECT DISTINCT a.codigo, a.nombre FROM asignaturas a
                JOIN tributaciones ar ON ar.asignatura_id = a.id
                WHERE ar.ra_id = ?
                ORDER BY a.semestre, a.codigo
            """, (ra_id,)).fetchall()
            txt_asig = "\n".join(f"• {a[0]} {a[1]}" for a in asigs) or "(ninguna)"
            _para(c_asig, txt_asig, size=9)

        doc.add_paragraph()

    conn.close()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if salida is None:
        salida = str(OUT_DIR / f"mapa_ra_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
    doc.save(salida)
    print(f"✓ Mapa de R.A. generado: {salida}")
    return salida


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "mapa":
        generar_mapa_progreso()
    elif len(sys.argv) > 1 and sys.argv[1] == "mapa_ra":
        generar_mapa_ra()
    else:
        print("Uso: python3 src/generador_word.py [mapa|mapa_ra]")