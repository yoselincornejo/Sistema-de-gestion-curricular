from docx import Document
import re

archivo = "/tmp/diag.docx"
doc = Document(archivo)

def limpiar(t):
        return re.sub(r"\s+", " ", t).strip() if t else ""

print(f"Archivo: {archivo}")
print(f"Total de tablas: {len(doc.tables)}\n")

for i, tabla in enumerate(doc.tables):
        filas = len(tabla.rows)
        cols = len(tabla.rows[0].cells) if filas > 0 else 0
        print(f"=== Tabla {i} === ({filas} filas × {cols} columnas)")

        # Mostrar primera fila completa (probablemente header)
        if filas > 0:
            primera = [limpiar(c.text)[:40] for c in tabla.rows[0].cells]
            print(f"  Header: {primera}")

        # Si tiene más de 1 fila, mostrar segunda fila (probable primer dato)
        if filas > 1:
            segunda = [limpiar(c.text)[:40] for c in tabla.rows[1].cells]
            print(f"  Fila 1: {segunda}")
        print()