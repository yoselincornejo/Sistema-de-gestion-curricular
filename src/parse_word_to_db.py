"""
parse_word_to_db.py — Extrae datos de los .docx e inserta en sistema.db.

Enfoque heurístico:
  1. Carga desde la BD el diccionario de códigos válidos (competencias,
     niveles_dominio, resultados_aprendizaje).
  2. Extrae TODO el texto del documento (párrafos + cada celda de cada tabla).
  3. Usa regex para detectar menciones a esos códigos → tributaciones.
  4. Cuarentena: si un archivo no produce código de asignatura ni tributaciones,
     se escribe en 'revision_manual.txt' y NO se toca la BD.

Uso:
    python3 src/parse_word_to_db.py            # procesa todos los .docx
    python3 src/parse_word_to_db.py --dry-run  # solo muestra; no modifica BD
    python3 src/parse_word_to_db.py --file "data/programas/Semestre 1/MAT 111.docx"
    python3 src/parse_word_to_db.py --verbose  # muestra cada RA encontrado

La BD debe existir (ejecutar init_db.py primero).
"""

import logging
import re
import sqlite3
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn as _qn

_MC_FALLBACK = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"
_W_T = _qn('w:t')

def _elem_text(elem) -> str:
    """Get text from an XML element, skipping mc:Fallback to avoid AlternateContent doubling."""
    parts = []
    _collect_text(elem, parts)
    return ''.join(parts)

def _collect_text(node, parts):
    if node.tag == _MC_FALLBACK:
        return
    if node.tag == _W_T and node.text:
        parts.append(node.text)
    for child in node:
        _collect_text(child, parts)

# ── Configuración ─────────────────────────────────────────────────────────────

RUTA_DB          = Path("data/sistema.db")
CARPETA_DOCS     = Path("data/programas")
LOG_REVISION     = Path("data/output/revision_manual.txt")
LOG_PARSE        = Path("data/output/parse_word_to_db.log")
# QUI 121 Lab: documento auxiliar sin estructura de programa (1 tabla).
EXCLUIDOS = {"QUI 121 Química para Ingeniería - Laboratorios"}

Path("data/output").mkdir(parents=True, exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format="%(levelname)-8s %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(LOG_PARSE, mode="w", encoding="utf-8"),
    ],
)
log = logging.getLogger(__name__)
VERBOSE = "--verbose" in sys.argv


# ══════════════════════════════════════════════════════════════════════════════
#  PASO 1 — Diccionario de códigos válidos desde la BD
# ══════════════════════════════════════════════════════════════════════════════

class DiccionarioCodigos:
    """
    Carga desde sistema.db todos los códigos válidos y construye el índice
    de búsqueda: texto_en_docx → ra_id en la BD.
    """

    def __init__(self, conn: sqlite3.Connection):
        # Competencias válidas: {'CL1', 'CE2', 'CG4', ...}
        self.competencias: set[str] = {
            r[0] for r in conn.execute("SELECT codigo FROM competencias")
        }

        # Niveles de dominio válidos por competencia:
        # {'CE1': {'ND1', 'ND2', 'ND3'}, ...}
        self.niveles: dict[str, set[str]] = {}
        for comp_cod, nd_cod in conn.execute(
            """SELECT c.codigo, nd.codigo_nivel
               FROM niveles_dominio nd JOIN competencias c ON c.id = nd.competencia_id"""
        ):
            self.niveles.setdefault(comp_cod, set()).add(nd_cod)

        # Mapa completo: codigo_completo → id
        # Incluye variantes con N-prefix y RA. para máxima compatibilidad.
        # ORDER BY codigo_nivel ASC garantiza que ND1 se procese primero;
        # las variantes "sin nivel" usan setdefault para que ND1 gane
        # cuando el documento omite el nivel de dominio.
        self.ra_por_codigo: dict[str, int] = {}
        for ra_id, nd_id, cod_ra, cod_completo in conn.execute(
            """SELECT ra.id, ra.nivel_dominio_id, ra.codigo_ra, ra.codigo_completo
               FROM resultados_aprendizaje ra
               JOIN niveles_dominio nd ON nd.id = ra.nivel_dominio_id
               ORDER BY nd.codigo_nivel ASC, ra.codigo_ra ASC"""
        ):
            # Forma canónica: "CE1, ND2, RA3"
            self.ra_por_codigo[_normalizar(cod_completo)] = ra_id

            # Descomponer para generar variantes
            m = re.match(r"([A-Z]{1,3}\d+),\s*(ND\d+),\s*(.+)", cod_completo)
            if not m:
                continue
            comp, nd, ra = m.group(1), m.group(2), m.group(3)
            num_nd = nd[2:]  # "ND2" → "2"

            # Variantes N-prefix: "CE1, N2, RA3"
            self.ra_por_codigo[_normalizar(f"{comp}, N{num_nd}, {ra}")] = ra_id
            # Sin nivel: "CE1, RA3" / "CE1, D1" → setdefault para que ND1 gane
            self.ra_por_codigo.setdefault(_normalizar(f"{comp}, {ra}"), ra_id)
            # Con punto: "CE1, ND2, RA.3"
            if ra.startswith("RA"):
                num = ra[2:]
                self.ra_por_codigo[_normalizar(f"{comp}, {nd}, RA.{num}")] = ra_id
                self.ra_por_codigo[_normalizar(f"{comp}, N{num_nd}, RA.{num}")] = ra_id
                self.ra_por_codigo.setdefault(_normalizar(f"{comp}, RA.{num}"), ra_id)
            # DC-codes: algunos docs usan "DC1" en vez de "D1"
            if ra.startswith("D"):
                num = ra[1:]
                self.ra_por_codigo[_normalizar(f"{comp}, {nd}, DC{num}")] = ra_id
                self.ra_por_codigo[_normalizar(f"{comp}, N{num_nd}, DC{num}")] = ra_id
                self.ra_por_codigo.setdefault(_normalizar(f"{comp}, DC{num}"), ra_id)

        # Patron regex que detecta CUALQUIER mención a un código tipo RA.
        # Soporta separadores coma O guión: "CL1, N2, RA1" y "CG2-ND2-DC1".
        comp_alternativa = "|".join(sorted(self.competencias, key=len, reverse=True))
        SEP = r"[\s]*[-,][\s]*"   # separador flexible: coma o guión con espacios opcionales
        self._PATRON_RA = re.compile(
            rf"(?P<comp>{comp_alternativa})"        # competencia: CL1, CE2, CG4…
            rf"(?:{SEP}"                            # separador opcional
            r"(?P<nivel>ND?\s*\d+))?"              # nivel: ND2 / N2  (nunca D/DC aquí)
            rf"{SEP}"                               # separador obligatorio
            r"(?P<ra>RA\.?\s*\d+|DC?\s*\d+)",      # RA1 / RA.1 / D1 / DC1
            re.IGNORECASE,
        )

    def extraer_ra_ids_de_texto(self, texto: str) -> set[int]:
        """
        Aplica el patrón sobre un bloque de texto.
        Reconstruye la clave canónica desde los grupos nombrados, por lo que
        los separadores guión/coma son transparentes para la búsqueda.
        """
        # Normalizar variantes como "CG 1 SUV" → "CG1" y "CG4 SUV" → "CG4":
        # 1) colapsar espacio entre letras y número: "CG 1" → "CG1"
        texto = re.sub(r"\b(C[EGL][A-Z]?)\s+(\d)\b", r"\1\2", texto, flags=re.IGNORECASE)
        # 2) quitar sufijo " SUV" que ya queda pegado al código: "CG1 SUV" → "CG1"
        texto = re.sub(r"\b(C[EGL]\d)\s+SUV\b", r"\1", texto, flags=re.IGNORECASE)
        encontrados = set()
        for m in self._PATRON_RA.finditer(texto):
            comp  = m.group("comp").upper().strip()
            nivel = (m.group("nivel") or "").upper().replace(" ", "")
            ra    = m.group("ra").upper().replace(" ", "").replace(".", "")

            # N2 → ND2
            if nivel and nivel.startswith("N") and not nivel.startswith("ND"):
                nivel = "ND" + nivel[1:]

            # DC1 → D1  (alias de D en documentos TIPE)
            if ra.startswith("DC"):
                ra = "D" + ra[2:]

            # Competencias CG con RA-codes → convertir a D-code (RA1 → D1)
            if comp.startswith("CG") and ra.startswith("RA"):
                ra = "D" + ra[2:]

            # Si el documento especifica nivel, solo buscar con ese nivel exacto.
            # No hacer fallback a otro nivel: si CG4-ND2-D3 no existe, no mapear
            # silenciosamente a CG4-ND1-D3 (sería una tributación incorrecta).
            claves = (
                [f"{comp}, {nivel}, {ra}"] if nivel
                else [f"{comp}, {ra}"]
            )
            ra_id = next((self.ra_por_codigo.get(c) for c in claves
                          if self.ra_por_codigo.get(c)), None)

            if ra_id:
                encontrados.add(ra_id)
            elif VERBOSE:
                log.debug("  sin match: %r → claves intentadas: %s",
                          m.group(0), claves)
        return encontrados


def _normalizar(s: str) -> str:
    """Quita espacios extras, puntos finales, y pasa a mayúsculas para comparar."""
    s = re.sub(r"\s+", " ", s).strip().rstrip(".")
    # Quitar espacio interior en "RA. 1" → "RA.1", "D C1" → "DC1"
    s = re.sub(r"(RA|DC?)\s*\.\s*(\d)", r"\1.\2", s, flags=re.IGNORECASE)
    return s.upper()


# ══════════════════════════════════════════════════════════════════════════════
#  PASO 2 — Extracción de texto completo del documento
# ══════════════════════════════════════════════════════════════════════════════

def extraer_texto_completo(doc: Document) -> str:
    """
    Concatena TODOS los párrafos y TODAS las celdas de TODAS las tablas.
    Se usa para la búsqueda heurística de códigos RA.
    """
    partes = []

    # Párrafos directos del documento
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            partes.append(t)

    # Cada celda de cada tabla (sin importar profundidad ni estructura)
    # Excluir tablas de unidades de aprendizaje: sus celdas mezclan RAs con
    # contenidos de unidades y pueden incluir desempeños extra no declarados
    # formalmente en la sección de resultados de aprendizaje.
    _HEADER_UNIDADES = {"unidades de aprendizaje", "contenidos"}
    for tabla in doc.tables:
        if tabla.rows:
            header_cells = {c.text.strip().lower() for c in tabla.rows[0].cells}
            if any(any(k in h for k in _HEADER_UNIDADES) for h in header_cells):
                continue
        for fila in tabla.rows:
            for celda in fila.cells:
                t = celda.text.strip()
                if t:
                    partes.append(t)

    return "\n".join(partes)


# ══════════════════════════════════════════════════════════════════════════════
#  PASO 3 — Extracción de metadatos principales
# ══════════════════════════════════════════════════════════════════════════════

# Patrones de código de asignatura: "MAT 111", "IMAT211", "CFG 111", "PRO 121"
_PATRON_CODIGO_ASIG = re.compile(
    r"\b([A-Z]{2,5}\s*\d{3}[A-Z]?)\b"
)

# Palabras clave que etiquetan el código en el documento
_PATRON_ETIQUETA_CODIGO = re.compile(
    r"(?:código[s]?|code|asignatura|curso)\s*[:\-]?\s*([A-Z]{2,5}\s*\d{3}[A-Z]?)",
    re.IGNORECASE,
)

def _extraer_codigo_y_nombre(doc: Document, texto_completo: str) -> tuple[str, str]:
    """
    Busca el código y nombre de la asignatura.
    Estrategia:
      1. Tabla 0, fila 1 → columnas típicas "Nombre / Código"
      2. Tabla con 2 cols y cabecera "Nombre"
      3. Búsqueda de etiquetas textuales ("Código: MAT 111")
      4. Primera mención de código alfanumérico en el texto
    """
    codigo = ""
    nombre = ""

    # Estrategia 1: Tabla 0 estructura estándar (8 filas × 6 cols)
    if doc.tables:
        t0 = doc.tables[0]
        try:
            fila_nombre = t0.rows[1]
            cs = [c.text.strip() for c in fila_nombre.cells]
            if len(cs) >= 5:
                cand_nombre = cs[1]
                cand_codigo = cs[4]
                if _PATRON_CODIGO_ASIG.match(cand_codigo):
                    nombre = cand_nombre
                    codigo = cand_codigo
        except IndexError:
            pass

    # Estrategia 2: tabla con cabecera "Nombre" en columna 0
    if not codigo:
        for t in doc.tables:
            rows = t.rows
            if len(rows) < 2 or len(t.columns) < 2:
                continue
            h = rows[0].cells[0].text.strip().lower()
            if "nombre" in h:
                try:
                    nombre_cand = rows[1].cells[1].text.strip()
                    # Buscar código en la fila siguiente
                    if len(rows) > 2:
                        cod_cand = rows[2].cells[1].text.strip()
                        if _PATRON_CODIGO_ASIG.match(cod_cand):
                            nombre = nombre_cand
                            codigo = cod_cand
                            break
                except IndexError:
                    pass

    # Estrategia 3: etiqueta textual
    if not codigo:
        m = _PATRON_ETIQUETA_CODIGO.search(texto_completo)
        if m:
            codigo = m.group(1).strip()

    # Estrategia 4: primera mención de código en el texto
    if not codigo:
        m = _PATRON_CODIGO_ASIG.search(texto_completo)
        if m:
            codigo = m.group(1).strip()

    # Normalizar código: colapsar espacios y asegurar espacio entre letras y números
    # "IMAT413" → "IMAT 413", "MAT  111" → "MAT 111"
    codigo = re.sub(r"\s+", " ", codigo).strip()
    codigo = re.sub(r"([A-Za-z])(\d)", r"\1 \2", codigo)

    return codigo, nombre


# ══════════════════════════════════════════════════════════════════════════════
#  PASO 4 — Extracción de metadatos secundarios
# ══════════════════════════════════════════════════════════════════════════════

def _txt_t(t, fila: int, col: int) -> str:
    try:
        return t.rows[fila].cells[col].text.strip()
    except IndexError:
        return ""

def _to_float(s: str) -> Optional[float]:
    s = re.sub(r"[^\d,\.]", "", s).replace(",", ".")
    try:
        return float(s)
    except (ValueError, TypeError):
        return None

def _to_int(s: str) -> Optional[int]:
    try:
        return int(re.sub(r"[^\d]", "", s))
    except (ValueError, TypeError):
        return None

def _semestre_desde_nivel(nivel: str) -> Optional[int]:
    if not nivel:
        return None
    n = nivel.strip()
    ROM = {"I":1,"II":2,"III":3,"IV":4,"V":5,"VI":6,"VII":7,
           "VIII":8,"IX":9,"X":10,"XI":11,"XII":12}

    # Numeral romano antes de "semestre": "III Semestre", "X Semestre del 3° Ciclo"
    m = re.search(r"\b(XII|XI|X|IX|VIII|VII|VI|V|IV|III|II|I)\s+[Ss]emestre", n)
    if m:
        return ROM.get(m.group(1).upper())

    # "Semestre N" (dígito después): "Semestre 3", "Semestre 9 de 3er ciclo"
    m = re.search(r"[Ss]emestre\s+(\d+)", n)
    if m:
        return int(m.group(1))

    # Dígito ordinal antes de "semestre": "3er Semestre", "4to semestre", "4º Semestre"
    m = re.search(r"(\d+)[°º\w.]*\s+[Ss]emestre", n, re.IGNORECASE)
    if m:
        return int(m.group(1))

    # Palabra ordinal: "Primer Semestre", "Cuarto Semestre"
    _ORD = {"primer":1,"segundo":2,"tercer":3,"cuarto":4,"quinto":5,
            "sexto":6,"séptimo":7,"octavo":8,"noveno":9,"décimo":10}
    for word, num in _ORD.items():
        if re.search(rf"\b{word}\b", n, re.IGNORECASE):
            return num

    # Año numérico como fallback: "1° Año" → S1, "2° Año" → S3 (primer sem del año)
    m_anio = re.search(r"(\d+)[°º]\s*[Aa]ño", n)
    if m_anio:
        return (int(m_anio.group(1)) - 1) * 2 + 1

    return None

def extraer_identificacion(doc: Document, texto_completo: str, codigo: str, nombre: str) -> dict:
    """Extrae metadatos numéricos y de texto de la tabla de identificación."""
    data = {
        "codigo": codigo, "nombre": nombre,
        "facultad": "", "carrera": "", "nivel": "",
        "duracion": "", "requisitos": "", "semestre": None,
        "horas_directa": None, "horas_autonoma": None,
        "semanas": None, "creditos": None,
    }

    if not doc.tables:
        return data

    t0 = doc.tables[0]
    try:
        data["facultad"]  = _txt_t(t0, 0, 1)
        data["carrera"]   = _txt_t(t0, 0, 4)
        if not nombre:
            data["nombre"] = _txt_t(t0, 1, 1)
        data["nivel"]     = _txt_t(t0, 2, 1)
        data["duracion"]  = _txt_t(t0, 2, 4)
        data["requisitos"]= _txt_t(t0, 3, 1)
        data["semestre"]  = _semestre_desde_nivel(data["nivel"])

        # Fila de horas → buscar de abajo hacia arriba la primera fila
        # cuya primera celda tenga un valor numérico (maneja tablas con
        # fila extra de descripción al final, como IMAT 413).
        fila_vals = None
        for row in reversed(t0.rows):
            vals = [c.text.strip() for c in row.cells]
            if len(vals) >= 6 and _to_float(vals[0]) is not None:
                fila_vals = vals
                break
        if fila_vals:
            data["horas_directa"]  = _to_float(fila_vals[0])
            data["horas_autonoma"] = _to_float(fila_vals[1])
            data["semanas"]        = _to_int(fila_vals[3])
            data["creditos"]       = _to_int(fila_vals[5])
    except (IndexError, AttributeError):
        pass

    return data


def extraer_descripcion(doc: Document) -> str:
    """
    Extrae la descripción de la asignatura. Busca en tres estrategias:
    1. Tabla 1-col que empieza con "La asignatura…"
    2. Celda dentro de tabla multi-col que contiene "DESCRIPCIÓN DE LA ASIGNATURA"
    3. Párrafos ubicados tras el encabezado "DESCRIPCIÓN DE LA ASIGNATURA:"
    """
    partes = []

    # Estrategia 1: tabla de 1 col que empieza con "La asignatura…"
    # "Esta asignatura aporta…" es la sección de Aporte al Perfil, NO la descripción.
    for t in doc.tables:
        if len(t.columns) == 1:
            txt = t.rows[0].cells[0].text.strip()
            if txt.lower().startswith("la asignatura"):
                for row in t.rows:
                    p = row.cells[0].text.strip()
                    if p:
                        partes.append(p)
    if partes:
        return "\n".join(partes)

    # Estrategia 2: celda en tabla multi-col que contiene "DESCRIPCIÓN DE LA ASIGNATURA"
    # (patrón: tabla de identificación donde la última fila tiene la descripción fusionada)
    for t in doc.tables:
        for row in t.rows:
            cell_txt = row.cells[0].text.strip()
            if re.search(r"DESCRIPCI[OÓ]N DE LA ASIGNATURA", cell_txt, re.IGNORECASE):
                # Saltar el encabezado "DESCRIPCIÓN DE LA ASIGNATURA:" y tomar el resto
                resto = re.split(r"DESCRIPCI[OÓ]N DE LA ASIGNATURA\s*[:\n]+", cell_txt,
                                 flags=re.IGNORECASE, maxsplit=1)
                if len(resto) > 1 and resto[1].strip():
                    return resto[1].strip()

    # Estrategia 3: párrafos (o tabla 1-col inmediata) tras "DESCRIPCIÓN DE LA ASIGNATURA:"
    # El encabezado puede estar como párrafo libre O dentro de una celda de tabla
    # (ej. IMAT 223: heading dentro de la tabla de identificación, descripción en párrafo siguiente)
    _STOP = {"aporte al perfil", "resultados de aprendizaje",
             "programa de la asignatura", "unidades de aprendizaje",
             "identificaci"}
    from docx.oxml.ns import qn as _qn

    def _tbl_has_desc_heading(tbl_elem):
        for tc in tbl_elem.iter(_qn('w:tc')):
            ct = "".join(r.text for r in tc.iter(_qn('w:t'))).strip().lower()
            if "descripci" in ct and "asignatura" in ct:
                return True
        return False

    capturing = False
    for child in doc.element.body:
        tag = child.tag.split('}')[-1]
        if tag == 'tbl':
            if capturing and not partes:
                # La descripción está en una tabla 1-col inmediatamente tras el encabezado
                # (ej. IMAT 421: tabla 1×1 con "Este es un curso de pregrado…")
                seen = set()
                cell_texts = []
                for tc in child.iter(_qn('w:tc')):
                    t = "".join(r.text for r in tc.iter(_qn('w:t'))).strip()
                    if t and t not in seen:
                        seen.add(t)
                        cell_texts.append(t)
                tbl_txt = "\n".join(cell_texts).strip()
                if tbl_txt and not any(k in tbl_txt.lower() for k in _STOP):
                    partes.append(tbl_txt)
                break
            if capturing:
                break
            # El encabezado puede estar dentro de esta tabla (ej. IMAT 223)
            if _tbl_has_desc_heading(child):
                capturing = True
            continue
        if tag != 'p':
            if capturing:
                break
            continue
        txt = _elem_text(child).strip()
        if not txt:
            continue
        tl = txt.lower()
        if "descripci" in tl and "asignatura" in tl and (tl.endswith(":") or tl.endswith("asignatura:")):
            # El párrafo puede contener la descripción + el heading en el mismo elemento
            # (ej. IMAT 223: AlternateContent con texto de descripción seguido de "2. DESCRIPCIÓN...")
            before = re.split(r"DESCRIPCI[OÓ]N DE LA ASIGNATURA", txt, flags=re.IGNORECASE, maxsplit=1)[0].strip()
            if before and len(before) > 50 and not any(k in before.lower() for k in _STOP):
                partes.append(before)
                break
            capturing = True
            continue
        if capturing:
            if any(k in tl for k in _STOP):
                break
            partes.append(txt)

    return "\n".join(partes)


def extraer_unidades(doc: Document) -> list[dict]:
    unidades = []
    for t in doc.tables:
        if len(t.columns) < 2:
            continue
        h0 = t.rows[0].cells[0].text.strip().lower()
        h1 = t.rows[0].cells[1].text.strip().lower() if len(t.columns) > 1 else ""
        if not ("resultado" in h0 or "desempeño" in h0 or "ra" == h0[:2] or
                "unidad" in h1 or "contenido" in h1):
            continue
        # Skip content-detail tables whose h0 is already the unit/workshop title
        if re.match(r"^(unidad|taller)\s+\d+", h0, re.IGNORECASE):
            continue
        for i, row in enumerate(t.rows[1:], 1):
            cells = row.cells
            # Col 0: RA/desempeños por unidad (texto de los RAs que aplican)
            desempenos  = cells[0].text.strip() if len(cells) > 0 else ""
            # Col 1: nombre de unidad + contenidos
            contenidos  = cells[1].text.strip() if len(cells) > 1 else ""
            # Col 2: indicador de logro (si existe en el documento)
            indicador   = cells[2].text.strip() if len(cells) > 2 else ""
            if not contenidos:
                continue
            # Si col 2 está vacía pero col 0 tiene texto de RA/desempeños, usarlo como indicador
            if not indicador and desempenos:
                dl = desempenos.lower()
                if not (dl.startswith("resultado") or dl.startswith("unidad")):
                    indicador = desempenos

            m = re.search(r"Unidad\s+(\d+|[IVX]+)", contenidos, re.IGNORECASE)
            num = None
            if m:
                raw = m.group(1)
                num = int(raw) if raw.isdigit() else \
                    {"I":1,"II":2,"III":3,"IV":4,"V":5,"VI":6,"VII":7,"VIII":8}.get(raw.upper())

            # Nombre de la unidad = primera línea no vacía del bloque de contenidos
            lines = [l.strip() for l in contenidos.split("\n") if l.strip()]
            nombre = lines[0][:120] if lines else contenidos[:120]

            unidades.append({
                "orden": num or i,
                "nombre": nombre,
                "contenidos": contenidos,
                "indicador_logro": indicador,
            })
    return unidades


_BIBLIO_HEADERS = {"autor", "título", "titulo", "editorial", "año", "anio",
                   "isbn", "ejemplares", "n°", "nº", "nro", "número", "numero",
                   "disponible", "biblioteca"}

def _parsear_biblio_tabla(t) -> list[dict]:
    entradas = []
    for row in t.rows[1:]:
        cs = [c.text.strip() for c in row.cells]
        # Saltar filas que parecen encabezados de columna
        non_empty = [c.lower().rstrip('.') for c in cs if c]
        if non_empty and sum(1 for c in non_empty if c in _BIBLIO_HEADERS) >= max(1, len(non_empty) // 2):
            continue
        if len(cs) >= 7:
            n, autor, titulo, editorial, anio, isbn, ej = cs[:7]
        elif len(cs) >= 4:
            n, autor, titulo, editorial = cs[:4]
            anio = isbn = ej = ""
        elif len(cs) >= 2:
            n, autor, titulo = "", cs[0], cs[1]
            editorial = anio = isbn = ej = ""
        else:
            continue
        if titulo or autor:
            entradas.append({"numero":n,"autor":autor,"titulo":titulo,
                             "editorial":editorial,"anio":anio,"isbn":isbn,"ejemplares":ej})
    return entradas


def extraer_bibliografia(doc: Document) -> tuple[list[dict], list[dict]]:
    basica, compl = [], []
    for t in doc.tables:
        h = t.rows[0].cells[0].text.strip().lower()
        if "básica" in h or "basica" in h:
            basica.extend(_parsear_biblio_tabla(t))
        elif "complementaria" in h:
            compl.extend(_parsear_biblio_tabla(t))
    return basica, compl


def extraer_linkografia(doc: Document) -> list[dict]:
    links = []
    for t in doc.tables:
        h = t.rows[0].cells[0].text.strip().lower()
        if "tipo de documento" not in h and "linkograf" not in h:
            continue
        for row in t.rows[1:]:
            cs = [c.text.strip() for c in row.cells]
            # Tabla estándar UV 8 cols:
            # tipo(0) autor(1) titulo_articulo(2) año(3) titulo_revista(4) vol(5) url(6) disponible(7)
            if len(cs) >= 7:
                links.append({
                    "tipo_documento":  cs[0],
                    "autor":           cs[1],
                    "titulo_articulo": cs[2],
                    "anio":            cs[3],
                    "titulo_revista":  cs[4],
                    "volumen":         cs[5] if len(cs) > 5 else "",
                    "url":             cs[6],
                    "disponible_en":   cs[7] if len(cs) > 7 else "",
                })
            elif len(cs) >= 3:
                links.append({
                    "tipo_documento": cs[0], "autor": "",
                    "titulo_articulo": cs[2] if len(cs) > 2 else "",
                    "anio": "", "titulo_revista": "",
                    "volumen": "", "url": cs[4] if len(cs) > 4 else "",
                    "disponible_en": "",
                })
            elif cs:
                for url in re.findall(r"https?://\S+", " ".join(cs)):
                    links.append({
                        "tipo_documento": "", "autor": "",
                        "titulo_articulo": "", "anio": "",
                        "titulo_revista": "", "volumen": "",
                        "url": url, "disponible_en": "",
                    })
    # URLs en tablas de recursos libres (1 col)
    for t in doc.tables:
        if len(t.columns) != 1:
            continue
        txt = " ".join(c.text for row in t.rows for c in row.cells)
        if re.search(r"https?://", txt):
            for url in re.findall(r"https?://\S+", txt):
                if not any(l["url"] == url for l in links):
                    links.append({"url": url, "titulo_articulo": "", "descripcion": "recursos"})
    return links


def extraer_otros_recursos(doc: Document) -> str:
    # Buscar el párrafo etiquetado "OTROS RECURSOS" y luego la primera tabla
    # que le sigue en el documento (usando el XML del body para preservar orden).
    from docx.oxml.ns import qn
    body = doc.element.body
    found_label = False
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            text = "".join(r.text or "" for r in child.iter(qn("w:t"))).strip().lower()
            if "otro" in text and "recurso" in text:
                found_label = True
        elif tag == "tbl" and found_label:
            # Primera tabla tras la etiqueta: extraer todo su texto
            from docx.table import Table
            tbl = Table(child, doc)
            lines = []
            for row in tbl.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        lines.append(txt)
            if lines:
                return "\n".join(lines)
            found_label = False  # tabla vacía, seguir buscando
    return ""


def extraer_responsables(doc: Document) -> dict:
    resp = {}
    for t in doc.tables:
        h = t.rows[0].cells[0].text.strip().lower()
        if "responsable" not in h:
            continue
        for row in t.rows:
            cs = [c.text.strip() for c in row.cells]
            if len(cs) >= 2:
                clave = cs[0].lower()
                valor = cs[1]
                if "docente" in clave:
                    resp["docente_a_cargo"] = valor
                elif "responsable" in clave:
                    resp["responsable"] = valor
                elif "versión" in clave or "version" in clave:
                    resp["version"] = valor
    return resp


_METOD_KEYS = {
    "aprendizaje", "estudio de", "clase expositiva", "taller",
    "simulac", "tutor", "terreno", "juego de", "servicio",
    "invertido", "colaborat", "cooperat", "problema", "proyecto",
    "pausa", "portafolio", "seminario", "exposit",
}

def extraer_metodologias(doc: Document) -> list[str]:
    metods = []
    for t in doc.tables:
        # Construir encabezado desde toda la primera fila (no solo la primera celda),
        # usando solo la primera línea de cada celda para evitar falsos positivos.
        header_cells = [c.text.strip().lower().split('\n')[0].strip() for c in t.rows[0].cells]
        h0_first = header_cells[0]
        header_all = " ".join(header_cells)

        # Excluir por tipo de encabezado conocido (cualquier celda de la fila)
        if any(k in header_all for k in ("resultado", "desempeño", "evaluaci", "tipo de eval",
                                         "experiencias de laboratorio", "tipo de documento",
                                         "bibliograf", "linkograf", "responsable",
                                         "facultad", "programa de la",
                                         "unidades de aprendizaje", "contenidos")):
            continue
        if any(k in h0_first for k in ("nombre", "unidad")):
            continue
        # Excluir tablas cuyo primer encabezado es prosa descriptiva de la asignatura
        _DESC_STARTERS = ("la asignatura", "esta asignatura", "al final de", "el programa",
                          "aporte al perfil", "esta carrera", "el presente",
                          "este curso", "este programa", "el curso")
        if any(h0_first.startswith(k) for k in _DESC_STARTERS):
            continue

        all_texts = [c.text.strip() for row in t.rows for c in row.cells]
        full_text = " ".join(all_texts)

        # Excluir tablas que son bibliografía/linkografía (contienen URLs)
        if re.search(r"https?://", full_text):
            continue

        # Verificar que la tabla contiene nombres de metodologías
        has_metod = any(
            any(k in txt.lower() for k in _METOD_KEYS)
            for txt in all_texts if len(txt) > 3
        )
        if not has_metod:
            continue

        ncols = len(t.columns)
        has_x = any(txt in ("X", "x", "✓") for txt in all_texts)

        if has_x and ncols in (2, 4):
            # Tabla de checkboxes: extraer solo las metodologías marcadas con X.
            # El layout puede ser (nombre, marca) o (marca, nombre) — se detecta
            # inspeccionando cuál columna del par es la más corta (la marca).
            def _par_nombre_marca(a, b):
                """Devuelve (nombre, marca) independientemente del orden en la tabla."""
                a_es_marca = len(a) <= 2 or a.upper() in ("X", "✓", "")
                b_es_marca = len(b) <= 2 or b.upper() in ("X", "✓", "")
                if a_es_marca and not b_es_marca:
                    return b, a
                return a, b

            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                pares_celdas = [(cells[0], cells[1])]
                if ncols >= 4:
                    pares_celdas.append((cells[2], cells[3]))
                for ca, cb in pares_celdas:
                    name, mark = _par_nombre_marca(ca, cb)
                    if name and mark.upper() in ("X", "✓") and len(name) > 3:
                        if name not in metods:
                            metods.append(name)
        else:
            # Tabla de texto libre: extraer todas las celdas con contenido
            for row in t.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt and len(txt) > 4 and txt not in metods:
                        metods.append(txt)
    return metods


_EVAL_HEADER_SKIP = {"porcentaje", "tipo de evaluaci", "evaluaciones formativas",
                     "evaluaciones sumativas", "tipo", "tipo de eval"}

def extraer_laboratorios(doc: Document) -> str:
    """Extrae la tabla '3.1 EXPERIENCIAS DE LABORATORIO' si existe."""
    for t in doc.tables:
        h0 = t.rows[0].cells[0].text.strip().lower()
        # Must be a dedicated lab section header (not just a table that mentions "laboratorio")
        if "experiencias de laboratorio" not in h0:
            continue
        lines = []
        for row in t.rows[1:]:
            cs = [_elem_text(c._tc).strip() for c in row.cells]
            unique = []
            seen = set()
            for c in cs:
                if c and c not in seen:
                    unique.append(c)
                    seen.add(c)
            if unique:
                lines.append(" | ".join(unique))
        if lines:
            return "\n".join(lines)
    return ""


def extraer_evaluaciones(doc: Document) -> list[dict]:
    evals = []
    for t in doc.tables:
        h = t.rows[0].cells[0].text.strip().lower()

        # Formato celda única: tabla 1×1 donde cada línea contiene "Tipo\t\tN%"
        if len(t.columns) == 1 and len(t.rows) == 1 and "%" in h:
            for line in t.rows[0].cells[0].text.strip().splitlines():
                parts = re.split(r'\t+', line.strip())
                if len(parts) >= 2:
                    tipo = parts[0].strip()
                    porc = parts[-1].strip()
                    if tipo and re.search(r'\d+\s*%', porc):
                        evals.append({"tipo": tipo, "porcentaje": porc})
            if evals:
                return evals

        # Accept table if first cell mentions evaluación (relaxed: no longer requires "tipo")
        if "evaluaci" not in h and "porcentaje" not in h:
            continue
        for row in t.rows[1:]:
            if len(row.cells) < 2:
                continue
            tc0 = row.cells[0]._tc
            tc1 = row.cells[1]._tc
            # Skip merged rows (description spanning both columns)
            if tc0 is tc1:
                continue
            tipo = _elem_text(tc0).strip()
            porc = _elem_text(tc1).strip()
            if not tipo:
                continue
            # Skip sub-header rows (e.g. "Evaluaciones sumativas" / "Porcentaje (%)")
            tipo_low = tipo.lower().rstrip('.')
            porc_low = porc.lower()
            if tipo_low in _EVAL_HEADER_SKIP or "porcentaje" in porc_low:
                continue
            # Skip rows that look like description text (very long tipo or starts with keywords)
            if tipo.startswith("Descripción") or tipo.startswith("Nota") or len(tipo) > 120:
                continue
            evals.append({"tipo": tipo, "porcentaje": porc})
    return evals


def extraer_descripcion_evaluaciones(doc: Document) -> str:
    """
    Captura la descripción de la sección de evaluaciones. Busca en dos lugares:
    1. Filas fusionadas (merged) dentro de la tabla de evaluaciones, que contienen
       texto descriptivo (Descripción general, política de notas, IA, etc.)
    2. Párrafos de texto libre que aparecen DESPUÉS de dicha tabla.
    """
    _STOP = {"recurso", "bibliograf", "linkograf", "otro recurso",
             "datos actual", "responsable"}
    capturing = False
    texts = []

    _WTR = _qn('w:tr')
    _WTC = _qn('w:tc')

    for child in doc.element.body:
        tag = child.tag.split('}')[-1]
        if tag == 'tbl':
            # Identificar tabla de evaluaciones por su PRIMERA FILA (no todo el texto)
            first_trs = child.findall(_WTR)
            if not first_trs:
                continue
            first_row_text = _elem_text(first_trs[0]).lower()
            if 'evaluaci' not in first_row_text and 'porcentaje' not in first_row_text:
                continue
            # Excluir tablas que no son de evaluaciones (linkografia, bibliografía…)
            if any(k in first_row_text for k in ("tipo de documento", "bibliograf",
                                                   "linkograf", "autor", "responsable")):
                continue
            # Es la tabla de evaluaciones: buscar filas fusionadas o de descripción
            for tr in child.findall(_WTR):
                tcs = tr.findall(_WTC)
                # Fila fusionada (gridSpan): solo 1 w:tc en el XML crudo
                if len(tcs) == 1:
                    txt = _elem_text(tcs[0]).strip()
                    if txt:
                        texts.append(txt)
                    continue
                if len(tcs) < 2:
                    continue
                # Fila con texto de descripción largo en col 0
                tipo = _elem_text(tcs[0]).strip()
                if (tipo.startswith("Descripción") or tipo.startswith("Nota")
                        or len(tipo) > 120):
                    texts.append(tipo)
            capturing = True
        elif tag == 'p' and capturing:
            txt = _elem_text(child).strip()
            if not txt:
                continue
            tl = txt.lower()
            if any(k in tl for k in _STOP):
                break
            texts.append(txt)

    return "\n".join(texts)


# ══════════════════════════════════════════════════════════════════════════════
#  PASO 5 — Parseo completo de un documento
# ══════════════════════════════════════════════════════════════════════════════

class ParseError(Exception):
    pass


def parsear_docx(ruta: Path, dicc: DiccionarioCodigos) -> dict:
    nombre = ruta.name

    try:
        doc = Document(str(ruta))
    except Exception as e:
        raise ParseError(f"No se pudo abrir el archivo: {e}")

    if len(doc.tables) < 2:
        raise ParseError(f"Documento no estándar: solo {len(doc.tables)} tabla(s)")

    # ── Texto completo (párrafos + celdas) ────────────────────────────────
    texto_completo = extraer_texto_completo(doc)

    # ── Código y nombre ───────────────────────────────────────────────────
    codigo, nombre_asig = _extraer_codigo_y_nombre(doc, texto_completo)

    if not codigo:
        raise ParseError("No se pudo determinar el código de la asignatura")

    # ── RAs heurísticos ───────────────────────────────────────────────────
    ra_ids = dicc.extraer_ra_ids_de_texto(texto_completo)

    if VERBOSE and ra_ids:
        conn_tmp = sqlite3.connect(str(RUTA_DB))
        for ra_id in sorted(ra_ids):
            cod = conn_tmp.execute(
                "SELECT codigo_completo FROM resultados_aprendizaje WHERE id=?", (ra_id,)
            ).fetchone()
            log.debug("  → RA encontrado: %s", cod[0] if cod else ra_id)
        conn_tmp.close()

    biblio_basica, biblio_compl = extraer_bibliografia(doc)
    return {
        "archivo":                  nombre,
        "identificacion":           extraer_identificacion(doc, texto_completo, codigo, nombre_asig),
        "descripcion":              extraer_descripcion(doc),
        "descripcion_evaluaciones": extraer_descripcion_evaluaciones(doc),
        "laboratorios":             extraer_laboratorios(doc),
        "ra_ids":                   ra_ids,
        "unidades":                 extraer_unidades(doc),
        "metodologias":             extraer_metodologias(doc),
        "evaluaciones":             extraer_evaluaciones(doc),
        "biblio_basica":            biblio_basica,
        "biblio_compl":             biblio_compl,
        "linkografia":              extraer_linkografia(doc),
        "otros_recursos":           extraer_otros_recursos(doc),
        "responsables":             extraer_responsables(doc),
    }


# ══════════════════════════════════════════════════════════════════════════════
#  PASO 6 — Inserción en BD
# ══════════════════════════════════════════════════════════════════════════════

def insertar_programa(conn: sqlite3.Connection, data: dict):
    ident  = data["identificacion"]
    codigo = ident["codigo"].strip()
    nombre = ident.get("nombre", "").strip()

    conn.execute("""
        INSERT INTO asignaturas
          (codigo, nombre, semestre, nivel, duracion, tipo, facultad, carrera,
           requisitos, horas_directa, horas_autonoma, semanas, creditos,
           descripcion, descripcion_evaluaciones, experiencias_laboratorio,
           otros_recursos, version, archivo_origen)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        ON CONFLICT(codigo) DO UPDATE SET
          nombre                   = excluded.nombre,
          semestre                 = excluded.semestre,
          nivel                    = excluded.nivel,
          duracion                 = excluded.duracion,
          facultad                 = excluded.facultad,
          carrera                  = excluded.carrera,
          requisitos               = excluded.requisitos,
          horas_directa            = excluded.horas_directa,
          horas_autonoma           = excluded.horas_autonoma,
          semanas                  = excluded.semanas,
          creditos                 = excluded.creditos,
          descripcion              = excluded.descripcion,
          descripcion_evaluaciones = excluded.descripcion_evaluaciones,
          experiencias_laboratorio = excluded.experiencias_laboratorio,
          otros_recursos           = excluded.otros_recursos,
          version                  = excluded.version,
          archivo_origen           = excluded.archivo_origen
    """, (
        codigo, nombre,
        ident.get("semestre"),       ident.get("nivel", ""),
        ident.get("duracion", ""),   "disciplinar",
        ident.get("facultad", ""),   ident.get("carrera", ""),
        ident.get("requisitos", ""), ident.get("horas_directa"),
        ident.get("horas_autonoma"), ident.get("semanas"),
        ident.get("creditos"),       data.get("descripcion", ""),
        data.get("descripcion_evaluaciones", ""),
        data.get("laboratorios", ""),
        data.get("otros_recursos",""),
        data.get("responsables", {}).get("version", ""),
        data.get("archivo", ""),
    ))

    asig_id = conn.execute(
        "SELECT id FROM asignaturas WHERE codigo = ?", (codigo,)
    ).fetchone()[0]

    # Limpiar datos previos (idempotente)
    for tbl in ("responsables","unidades","metodologias","evaluaciones",
                "bibliografia","linkografia","tributaciones"):
        conn.execute(f"DELETE FROM {tbl} WHERE asignatura_id = ?", (asig_id,))

    resp = data.get("responsables", {})
    for rol in ("responsable", "docente_a_cargo"):
        if resp.get(rol):
            conn.execute(
                "INSERT INTO responsables (asignatura_id, rol, nombre) VALUES (?,?,?)",
                (asig_id, rol, resp[rol])
            )

    for u in data.get("unidades", []):
        conn.execute("""INSERT INTO unidades
            (asignatura_id, orden, nombre, contenidos, indicador_logro) VALUES (?,?,?,?,?)""",
            (asig_id, u["orden"], u["nombre"], u["contenidos"], u["indicador_logro"]))

    for m in data.get("metodologias", []):
        if m:
            conn.execute(
                "INSERT INTO metodologias (asignatura_id, descripcion) VALUES (?,?)",
                (asig_id, m))

    for ev in data.get("evaluaciones", []):
        if ev.get("tipo"):
            conn.execute(
                "INSERT INTO evaluaciones (asignatura_id, tipo, porcentaje) VALUES (?,?,?)",
                (asig_id, ev["tipo"], ev.get("porcentaje", "")))

    for b in data.get("biblio_basica", []):
        conn.execute("""INSERT INTO bibliografia
            (asignatura_id,tipo,numero,autor,titulo,editorial,anio,isbn,ejemplares)
            VALUES (?,?,?,?,?,?,?,?,?)""",
            (asig_id,"basica",b.get("numero",""),b.get("autor",""),b.get("titulo",""),
             b.get("editorial",""),b.get("anio",""),b.get("isbn",""),b.get("ejemplares","")))

    for b in data.get("biblio_compl", []):
        conn.execute("""INSERT INTO bibliografia
            (asignatura_id,tipo,numero,autor,titulo,editorial,anio,isbn,ejemplares)
            VALUES (?,?,?,?,?,?,?,?,?)""",
            (asig_id,"complementaria",b.get("numero",""),b.get("autor",""),b.get("titulo",""),
             b.get("editorial",""),b.get("anio",""),b.get("isbn",""),b.get("ejemplares","")))

    for lk in data.get("linkografia", []):
        conn.execute("""INSERT INTO linkografia
            (asignatura_id, tipo_documento, autor, titulo_articulo, anio,
             titulo_revista, volumen, url, disponible_en)
            VALUES (?,?,?,?,?,?,?,?,?)""",
            (asig_id,
             lk.get("tipo_documento",""), lk.get("autor",""),
             lk.get("titulo_articulo",""), lk.get("anio",""),
             lk.get("titulo_revista",""), lk.get("volumen",""),
             lk.get("url",""), lk.get("disponible_en","")))

    # Tributaciones desde ra_ids ya validados
    for ra_id in data.get("ra_ids", set()):
        conn.execute("""INSERT OR IGNORE INTO tributaciones (asignatura_id, ra_id)
                        VALUES (?,?)""", (asig_id, ra_id))


# ══════════════════════════════════════════════════════════════════════════════
#  PASO 7 — Procesamiento masivo
# ══════════════════════════════════════════════════════════════════════════════

def procesar_todos(dry_run: bool = False, archivo_unico: Optional[Path] = None):
    if not RUTA_DB.exists():
        log.error("BD no encontrada: %s. Ejecuta init_db.py primero.", RUTA_DB)
        sys.exit(1)

    conn_dicc = sqlite3.connect(str(RUTA_DB))
    dicc = DiccionarioCodigos(conn_dicc)
    conn_dicc.close()

    log.info("Diccionario cargado: %d competencias | %d RA codes válidos",
             len(dicc.competencias), len(dicc.ra_por_codigo))

    archivos = [archivo_unico] if archivo_unico else sorted(CARPETA_DOCS.rglob("*.docx"))

    log.info("=" * 65)
    log.info("parse_word_to_db.py  —  %d documentos  %s",
             len(archivos), "[DRY-RUN]" if dry_run else "")
    log.info("=" * 65)

    conn = None if dry_run else sqlite3.connect(str(RUTA_DB))
    if conn:
        conn.execute("PRAGMA foreign_keys = ON")

    ok = errores = omitidos = cuarentena = 0
    en_cuarentena: list[tuple[str, str]] = []

    for ruta in archivos:
        nombre = ruta.name

        if any(ex in nombre for ex in EXCLUIDOS):
            log.info("OMITIDO   %s", nombre)
            omitidos += 1
            continue

        try:
            data = parsear_docx(ruta, dicc)
        except ParseError as e:
            log.error("ERROR     %-50s → %s", nombre, e)
            errores += 1
            en_cuarentena.append((nombre, str(e)))
            continue
        except Exception as e:
            log.exception("EXCEPCION %-45s → %s", nombre, e)
            errores += 1
            en_cuarentena.append((nombre, str(e)))
            continue

        codigo  = data["identificacion"]["codigo"]
        ra_ids  = data["ra_ids"]
        n_u     = len(data["unidades"])
        n_bb    = len(data["biblio_basica"])
        n_bc    = len(data["biblio_compl"])

        # ── Sistema de Cuarentena ─────────────────────────────────────────
        # Solo bloquear si no se puede identificar la asignatura.
        # Documentos sin tributaciones (electivos, idiomas con plan antiguo)
        # se cargan igual con sus metadatos.
        if not codigo:
            motivo = "código de asignatura no determinado"
            log.warning("CUARENTENA %-45s → %s", nombre, motivo)
            en_cuarentena.append((nombre, motivo))
            cuarentena += 1
            continue   # NO tocar la BD

        if not ra_ids:
            log.warning("SIN-RA    %-10s %-40s (se carga sin tributaciones)",
                        codigo, data["identificacion"].get("nombre","")[:39])

        # ── Inserción ─────────────────────────────────────────────────────
        estado = "dry-run" if dry_run else "ok"
        if not dry_run:
            insertar_programa(conn, data)

        log.info("%-10s %-10s %-40s | RA:%-3d U:%-3d BB:%-3d BC:%d",
                 estado.upper(), codigo,
                 data["identificacion"].get("nombre","")[:39],
                 len(ra_ids), n_u, n_bb, n_bc)
        ok += 1

    if conn:
        conn.commit()
        conn.close()

    # ── Archivo revision_manual.txt ───────────────────────────────────────
    if en_cuarentena:
        with open(LOG_REVISION, "w", encoding="utf-8") as f:
            f.write("Documentos que requieren revisión manual\n")
            f.write("=" * 60 + "\n\n")
            for nombre, motivo in en_cuarentena:
                f.write(f"{nombre}\n  Motivo: {motivo}\n\n")
        log.info("Archivo de revisión guardado en: %s", LOG_REVISION)

    log.info("")
    log.info("=" * 65)
    log.info("RESULTADO: %d OK | %d cuarentena | %d omitidos | %d errores",
             ok, cuarentena, omitidos, errores)
    if not dry_run:
        _verificar_bd()


def _verificar_bd():
    conn = sqlite3.connect(str(RUTA_DB))
    log.info("")
    log.info("ESTADO FINAL DE LA BD:")
    for tbl in ("asignaturas","tributaciones","unidades","bibliografia",
                "linkografia","metodologias","evaluaciones"):
        n = conn.execute(f"SELECT COUNT(*) FROM {tbl}").fetchone()[0]
        log.info("  %-25s %d", tbl, n)
    conn.close()


# ══════════════════════════════════════════════════════════════════════════════
#  Punto de entrada
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    dry_run       = "--dry-run" in sys.argv
    archivo_unico = None

    if "--file" in sys.argv:
        idx = sys.argv.index("--file")
        if idx + 1 < len(sys.argv):
            archivo_unico = Path(sys.argv[idx + 1])
            if not archivo_unico.exists():
                log.error("Archivo no encontrado: %s", archivo_unico)
                sys.exit(1)

    procesar_todos(dry_run=dry_run, archivo_unico=archivo_unico)
